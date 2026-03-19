#!/usr/bin/env python3
"""
MITRE ATT&CK 测试日志分析工具
自动解析goart生成的日志文件，生成安全验证报告
python3 analyze_atomic_logs2.py --format html
python3 analyze_atomic_logs2.py --format excel
python3 analyze_atomic_logs2.py --format docx
"""

import os
import re
import json
import glob
import pandas as pd
from datetime import datetime, timedelta
from collections import defaultdict, Counter
import argparse
import sys
import statistics

# 可选导入 docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class AtomicLogAnalyzer:
    def __init__(self, log_dir="./logs", report_dir="./reports"):
        self.log_dir = log_dir
        self.report_dir = report_dir
        self.results = []
        self.technique_stats = defaultdict(lambda: {
            "total": 0,
            "passed": 0,
            "failed": 0,
            "error": 0,
            "execution_times": [],
            "findings": [],
            "commands": [],
            "timestamps": [],
            "error_types": [],
            "files": []
        })

        # MITRE ATT&CK 战术映射
        self.tactic_mapping = {
            "T1": "初始访问", "T1078": "初始访问", "T1189": "初始访问", "T1190": "初始访问",
            "T1059": "执行", "T1064": "执行", "T1204": "执行", "T1203": "执行",
            "T1136": "持久化", "T1547": "持久化", "T1053": "持久化", "T1543": "持久化",
            "T1548": "权限提升", "T1068": "权限提升", "T1134": "权限提升", "T1055": "权限提升",
            "T1027": "防御规避", "T1112": "防御规避", "T1070": "防御规避", "T1036": "防御规避",
            "T1003": "凭据访问", "T1005": "凭据访问", "T1056": "凭据访问", "T1212": "凭据访问",
            "T1083": "发现", "T1018": "发现", "T1518": "发现", "T1057": "发现",
            "T1021": "横向移动", "T1534": "横向移动", "T1570": "横向移动",
            "T1041": "数据收集", "T1005": "数据收集", "T1560": "数据收集",
            "T1048": "数据外泄", "T1041": "数据外泄", "T1020": "数据外泄",
            "T1071": "命令与控制", "T1095": "命令与控制", "T1132": "命令与控制",
        }

        # 风险等级定义
        self.risk_levels = {
            "critical": ["T1003", "T1055", "T1059", "T1548"],
            "high": ["T1021", "T1048", "T1071", "T1136", "T1547", "T1053"],
            "medium": ["T1083", "T1018", "T1056", "T1057", "T1518", "T1036", "T1070"],
            "low": ["T1005", "T1041", "T1020", "T1132", "T1560", "T1534", "T1570"]
        }

        # 技术详细描述
        self.technique_descriptions = {
            "T1003": "操作系统凭据转储 - 从内存或SAM数据库中提取凭据",
            "T1055": "进程注入 - 将代码注入到其他进程内存中执行",
            "T1059": "命令行界面 - 使用命令行界面执行命令",
            "T1068": "利用漏洞提升权限 - 利用软件漏洞获取更高权限",
            "T1078": "有效账户 - 使用被盗用的有效账户进行访问",
            "T1134": "访问令牌操纵 - 操纵访问令牌以提升权限",
            "T1136": "创建账户 - 创建本地或域账户以维持访问",
            "T1547": "启动项持久化 - 通过启动项实现持久化",
            "T1548": "权限提升执行 - 利用系统特性提升权限",
            "T1021": "远程服务 - 使用远程服务进行横向移动",
            "T1048": "通过备用信道外泄数据 - 使用非标准协议外泄数据",
            "T1071": "应用层协议 - 使用应用层协议进行C2通信",
            "T1083": "文件和目录发现 - 枚举文件和目录信息",
            "T1005": "本地数据源 - 访问本地存储的数据",
            "T1053": "计划任务 - 创建或修改计划任务实现持久化",
            "T1056": "输入捕获 - 捕获用户输入以获取凭据",
            "T1057": "进程发现 - 获取进程信息以了解系统状态",
            "T1064": "脚本 - 使用脚本执行恶意代码",
            "T1070": "清除证据 - 删除或修改安全日志",
            "T1112": "修改注册表 - 修改注册表以实现持久化或配置更改",
            "T1189": "路过式入侵 - 通过访问网站植入恶意代码",
            "T1190": "利用面向公众的应用程序 - 利用Web应用程序漏洞",
            "T1203": "利用客户端执行 - 利用客户端应用程序漏洞",
            "T1204": "用户执行 - 诱导用户执行恶意代码",
            "T1518": "软件发现 - 枚举安装的软件",
            "T1020": "自动外泄 - 自动将数据传输到外部系统",
            "T1027": "混淆文件或信息 - 使用混淆技术隐藏恶意内容",
            "T1036": "伪装 - 伪装恶意软件以逃避检测",
            "T1041": "通过C2通道外泄 - 使用C2通道传输数据",
            "T1095": "非应用层协议 - 使用非标准协议进行通信",
            "T1132": "数据编码 - 对数据进行编码以逃避检测",
            "T1212": "利用凭据访问 - 利用凭据访问漏洞",
            "T1534": "横向移动工具 - 使用工具进行横向移动",
            "T1543": "创建或修改系统进程 - 创建系统级进程",
            "T1560": "收集归档数据 - 将数据归档以便外泄",
            "T1570": "横向移动传输工具 - 传输工具到目标系统",
            "T1001": "数据混淆 - 混淆数据以逃避检测",
            "T1002": "数据压缩 - 压缩数据以方便传输",
            "T1006": "直接卷访问 - 直接访问卷以绕过访问控制",
            "T1007": "系统服务发现 - 枚举系统服务",
            "T1008": " fallback 通道 - 使用备用C2通道",
            "T1010": "应用程序窗口发现 - 枚举应用程序窗口",
            "T1012": "查询注册表 - 查询注册表获取信息",
            "T1014": "Rootkit - 使用Rootkit隐藏存在",
            "T1016": "系统网络配置发现 - 获取网络配置信息",
            "T1018": "远程系统发现 - 发现网络上的其他系统",
            "T1022": "加密数据 - 加密数据以逃避检测",
            "T1024": "与用户通信 - 与目标用户建立通信",
            "T1025": "从可移动介质获取数据 - 访问可移动存储设备",
            "T1028": "Windows远程管理 - 使用WinRM进行远程管理",
            "T1029": "计划数据传输 - 按计划传输数据",
            "T1030": "连续传输数据 - 连续传输小块数据",
            "T1031": "修改现有服务 - 修改服务配置",
            "T1033": "系统所有者/用户发现 - 获取用户信息",
            "T1034": "路径拦截 - 利用路径优先级执行恶意代码",
            "T1035": "服务执行 - 通过服务执行恶意代码",
            "T1037": "登录脚本 - 使用登录脚本实现持久化",
            "T1038": "DLL搜索顺序劫持 - 劫持DLL加载顺序",
            "T1039": "网络共享发现 - 枚举网络共享",
            "T1040": "网络嗅探 - 捕获网络流量",
            "T1042": "更改默认文件关联 - 修改文件关联",
            "T1043": "伪装服务 - 创建伪装服务",
            "T1044": "文件系统权限弱点 - 利用文件权限弱点",
            "T1045": "软件打包 - 使用打包工具混淆代码",
            "T1046": "网络服务扫描 - 扫描网络服务",
            "T1047": "Windows管理规范 - 使用WMI执行操作",
            "T1049": "系统网络连接发现 - 获取网络连接信息",
            "T1050": "新服务 - 安装新服务",
            "T1051": "共享Webroot - 通过共享Webroot传播",
            "T1052": "外设设备访问 - 使用外设传输数据",
            "T1054": "修改系统镜像 - 修改系统镜像文件",
            "T1058": "服务停止 - 停止安全相关服务",
            "T1060": "注册表Run键 - 使用Run键实现持久化",
            "T1061": "图形界面 - 使用GUI执行操作",
            "T1062": "额外的Windows二进制文件 - 使用额外二进制文件",
            "T1063": "安全软件发现 - 枚举安全软件",
            "T1065": "非标准端口 - 使用非标准端口通信",
            "T1066": "混淆指示器移除 - 移除入侵痕迹",
            "T1067": "Bootkit - 使用Bootkit持久化",
            "T1069": "权限组发现 - 枚举权限组",
            "T1072": "第三方软件 - 使用第三方软件执行",
            "T1073": "DLL侧载 - 利用DLL侧载技术",
            "T1074": "数据分级 - 对收集的数据进行分级",
            "T1075": "传递哈希 - 使用哈希传递技术",
            "T1076": "远程桌面协议 - 使用RDP进行远程访问",
            "T1077": "Windows管理员共享 - 访问管理共享",
            "T1079": "多跳代理 - 使用多级代理隐藏来源",
            "T1080": "感染内容 - 通过感染内容传播",
            "T1081": "浏览器凭证提取 - 从浏览器提取凭据",
            "T1082": "系统信息发现 - 获取系统基本信息",
            "T1084": "Windows管理规范事件订阅 - 使用WMI事件",
            "T1085": " Rundll32 - 使用Rundll32执行代码",
            "T1086": "PowerShell - 使用PowerShell执行代码",
            "T1087": "账户发现 - 枚举用户账户",
            "T1088": "Bash历史 - 读取Bash历史记录",
            "T1089": "禁用安全工具 - 禁用安全软件",
            "T1090": "代理 - 使用代理服务器",
            "T1091": "通过可移动介质复制 - 通过可移动介质传播",
            "T1093": "进程空心化 - 使用进程空心化技术",
            "T1096": "NTFS文件属性 - 利用NTFS属性隐藏数据",
            "T1097": "传递票据 - 使用Kerberos票据传递",
            "T1098": "账户操纵 - 修改账户配置",
            "T1099": "修改文件时间 - 修改文件时间戳",
            "T1100": "Web Shell - 部署Web Shell",
            "T1101": "安全支持提供程序 - 安装恶意SSP",
            "T1102": "Web服务 - 使用Web服务进行C2",
            "T1103": "AppCert DLL - 使用AppCert DLL",
            "T1104": "多阶段通道 - 使用多级C2通道",
            "T1105": "远程文件复制 - 远程复制文件",
            "T1106": "本地API执行 - 通过本地API执行",
            "T1107": "文件删除 - 删除文件以清除证据",
            "T1108": "冗余访问 - 建立冗余访问路径",
            "T1109": "组件对象模型 - 使用COM执行",
            "T1110": "暴力破解 - 使用暴力破解凭据",
            "T1111": "多因素认证拦截 - 拦截MFA凭证",
            "T1113": "屏幕捕获 - 捕获屏幕内容",
            "T1114": "邮件收集 - 收集电子邮件数据",
            "T1115": "剪贴板数据 - 访问剪贴板内容",
            "T1116": "代码签名 - 使用代码签名伪装",
            "T1117": "Regsvr32 - 使用Regsvr32执行",
            "T1118": "InstallUtil - 使用InstallUtil执行",
            "T1119": "自动化收集 - 自动收集数据",
            "T1120": "外设发现 - 枚举外设设备",
            "T1121": "Regsvr32执行 - 使用Regsvr32",
            "T1122": "组件对象模型劫持 - 劫持COM对象",
            "T1123": "音频捕获 - 捕获音频内容",
            "T1124": "系统时间发现 - 获取系统时间信息",
            "T1125": "视频捕获 - 捕获视频内容",
            "T1126": "网络共享连接 - 连接网络共享",
            "T1127": "受信任开发者工具 - 使用受信任工具",
            "T1128": "Netsh助手DLL - 使用Netsh",
            "T1129": "执行通过模块加载 - 通过模块加载执行",
            "T1130": "Kerberoasting - 利用Kerberoasting",
            "T1131": "Netsh - 使用Netsh执行",
            "T1132": "数据编码 - 对数据进行编码",
            "T1133": "外部远程服务 - 使用外部服务",
            "T1135": "网络共享发现 - 发现网络共享",
            "T1137": "Office应用启动 - 通过Office启动",
            "T1138": "应用补丁绕过 - 绕过应用补丁",
            "T1140": "反混淆/解码文件 - 解码混淆文件",
            "T1141": "用户欺骗 - 欺骗用户执行",
            "T1142": "键盘记录 - 记录键盘输入",
            "T1143": "隐藏进程 - 隐藏恶意进程",
            "T1144": "Gatekeeper绕过 - 绕过Gatekeeper",
            "T1145": "SSH私钥发现 - 发现SSH私钥",
            "T1146": "清除命令历史 - 清除历史记录",
            "T1147": "隐藏文件 - 隐藏恶意文件",
            "T1148": "HISTCONTROL - 操纵历史控制",
            "T1149": "LC_PRELOAD - 使用LD_PRELOAD",
            "T1150": "Plist修改 - 修改Plist文件",
            "T1151": "Shadow Copies - 操作卷影副本",
            "T1152": "源端口伪造 - 伪造源端口",
            "T1153": "源路径伪造 - 伪造源路径",
            "T1154": "Trap命令 - 使用Trap命令",
            "T1155": "AppleScript - 使用AppleScript",
            "T1156": ".bash_profile修改 - 修改bash配置",
            "T1157": "Dylib劫持 - 劫持动态库",
            "T1158": "隐藏文件和目录 - 隐藏文件",
            "T1159": "LaunchAgent - 使用LaunchAgent",
            "T1160": "LaunchDaemon - 使用LaunchDaemon",
            "T1161": "LC_LOAD_DYLIB - 修改动态库加载",
            "T1162": "登录项 - 添加登录项",
            "T1163": "LaunchAgent执行 - 执行LaunchAgent",
            "T1164": "LaunchDaemon执行 - 执行LaunchDaemon",
            "T1165": "系统定位器 - 修改定位器",
            "T1166": "Setuid和Setgid - 利用特殊权限",
            "T1167": "动态链接器劫持 - 劫持动态链接器",
            "T1168": "本地作业调度 - 调度本地作业",
            "T1169": "Sudo缓存 - 利用Sudo缓存",
            "T1170": "Mshta - 使用Mshta执行",
            "T1171": "BITS作业 - 使用BITS",
            "T1172": "Rundll32执行 - 使用Rundll32",
            "T1173": "动态数据交换 - 使用DDE",
            "T1174": "密码过滤器DLL - 安装密码过滤器",
            "T1175": "分布式组件对象模型 - 使用DCOM",
            "T1176": "浏览器扩展 - 安装恶意扩展",
            "T1177": "组件对象模型劫持 - 劫持COM",
            "T1178": "WMI持久化 - 使用WMI持久化",
            "T1179": "Hook - 使用API Hook",
            "T1180": "Screensaver - 利用屏幕保护程序",
            "T1181": "DefaultAssoc - 修改默认关联",
            "T1182": "AppCert - 使用AppCert",
            "T1183": "AppInit DLL - 使用AppInit",
            "T1184": "Winlogon - 修改Winlogon",
            "T1185": "浏览器入口 - 操纵浏览器",
            "T1186": "ProcessDoppelganging - 使用进程替身",
            "T1187": "Forced认证 - 强制认证",
            "T1188": "多播接收 - 使用多播",
            "T1189": "路过式入侵 - 路过式攻击",
            "T1190": "利用Web应用 - 利用Web应用",
            "T1191": "CMSTP - 使用CMSTP",
            "T1192": "高信任认证 - 滥用高信任",
            "T1193": "执行通过用户 - 通过用户执行",
            "T1194": "执行通过硬件 - 通过硬件执行",
            "T1195": "供应链攻击 - 通过供应链进行攻击",
            "T1197": "BITS Jobs - 使用BITS进行持久化",
            "T1199": "受信任关系 - 滥用受信任关系",
            "T1200": "硬件添加 - 通过硬件添加持久化",
            "T1201": "密码策略发现 - 发现密码策略",
            "T1202": "间接命令执行 - 间接执行命令",
            "T1203": "利用客户端执行 - 利用客户端软件",
            "T1204": "用户执行 - 诱导用户执行恶意代码",
            "T1205": "网络流量信号 - 使用网络流量作为信号",
            "T1210": "利用远程服务 - 利用远程服务漏洞",
            "T1211": "利用防御逃逸 - 逃避防御检测",
            "T1212": "利用凭据访问 - 利用凭据访问漏洞",
            "T1213": "数据从信息库 - 从信息库收集数据",
            "T1216": "系统脚本 - 使用系统脚本",
            "T1217": "浏览器书签发现 - 发现浏览器书签",
            "T1218": "系统二进制代理执行 - 使用系统二进制文件",
            "T1219": "远程访问软件 - 使用远程访问软件",
            "T1220": "XSL脚本处理 - 使用XSL脚本",
            "T1221": "模板注入 - 进行模板注入攻击",
            "T1547": "启动项持久化 - 通过启动项实现持久化",
            "T1548": "权限提升执行 - 提升权限执行代码",
            "T1571": "非标准端口 - 通过非标准端口进行C2通信",
            "T1572": "协议隧道 - 使用协议隧道技术",
        }

        # 错误类型分类
        self.error_patterns = {
            "权限不足": r"permission denied|access denied|unauthorized|没有权限",
            "命令未找到": r"command not found|不是内部或外部命令|unknown command",
            "网络错误": r"connection refused|timeout|network|unreachable",
            "资源不足": r"no space|memory|resource|out of memory",
            "依赖缺失": r"module not found|import error|no module",
            "配置错误": r"config|configuration|设置错误",
            "语法错误": r"syntax|invalid syntax|语法错误"
        }

    def parse_log_file(self, log_file: str) -> dict:
        """解析日志文件，提取关键信息"""
        result = {
            "file": log_file,
            "filename": os.path.basename(log_file),
            "technique": None,
            "timestamp": None,
            "status": "unknown",
            "execution_time": None,
            "commands": [],
            "output": [],
            "errors": [],
            "warnings": [],
            "indicators": [],
            "test_name": None,
            "description": None,
            "tactic": "未知",
            "error_type": None,
            "output_size": 0,
            "line_count": 0
        }

        try:
            # 从文件名提取技术ID和时间戳
            filename = os.path.basename(log_file)

            # 提取技术ID (如 T1003)
            tech_match = re.search(r'(T\d{4})', filename)
            if tech_match:
                result["technique"] = tech_match.group(1)
                result["description"] = self.technique_descriptions.get(result["technique"], "未知技术")
                # 确定战术分类
                result["tactic"] = self.tactic_mapping.get(result["technique"], "未知")

            # 提取时间戳 (如 20260317_050314)
            time_match = re.search(r'(\d{8})_(\d{6})', filename)
            if time_match:
                date_str = time_match.group(1)
                time_str = time_match.group(2)
                result["timestamp"] = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]} {time_str[:2]}:{time_str[2:4]}:{time_str[4:6]}"

            # 读取日志内容
            with open(log_file, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                lines = content.split('\n')
                result["line_count"] = len(lines)
                result["output_size"] = len(content)

            # 提取关键信息
            for i, line in enumerate(lines):
                line = line.strip()

                # 提取状态
                if "Test completed successfully" in line or "SUCCESS" in line.upper():
                    result["status"] = "passed"
                elif "Test failed" in line or "FAILED" in line.upper():
                    result["status"] = "failed"
                elif "ERROR" in line.upper():
                    result["status"] = "error"
                    result["errors"].append(line)
                    # 分类错误类型
                    for error_name, pattern in self.error_patterns.items():
                        if re.search(pattern, line, re.IGNORECASE):
                            result["error_type"] = error_name
                            break
                elif "WARNING" in line.upper():
                    result["warnings"].append(line)

                # 提取执行命令
                if "Executing command:" in line or "Command:" in line or "Running:" in line:
                    cmd_match = re.search(r'(?:Executing command:|Command:|Running:)\s*(.+)', line)
                    if cmd_match:
                        result["commands"].append(cmd_match.group(1))

                # 提取执行时间
                if "execution time" in line.lower() or "duration" in line.lower():
                    time_match = re.search(r'(\d+\.?\d*)\s*(?:seconds|s)', line.lower())
                    if time_match:
                        result["execution_time"] = float(time_match.group(1))

                # 提取测试名称
                if "Test name:" in line or "Test:" in line:
                    name_match = re.search(r'(?:Test name:|Test:)\s*(.+)', line)
                    if name_match and not result["test_name"]:
                        result["test_name"] = name_match.group(1)

                # 提取检测指标
                if "Indicator:" in line or "Detected:" in line or "Finding:" in line:
                    indicator_match = re.search(r'(?:Indicator:|Detected:|Finding:)\s*(.+)', line)
                    if indicator_match:
                        result["indicators"].append(indicator_match.group(1))

                # 提取输出信息（最后10行作为输出预览）
                if i >= len(lines) - 10:
                    if line and len(line) < 500:  # 避免太长的行
                        result["output"].append(line)

            # 如果没有明确状态，根据内容推断
            if result["status"] == "unknown":
                if any(word in content.upper() for word in ["ERROR", "FAILED", "FAILURE"]):
                    result["status"] = "failed"
                elif any(word in content.upper() for word in ["SUCCESS", "COMPLETED", "PASSED"]):
                    result["status"] = "passed"
                elif "WARNING" in content.upper():
                    result["status"] = "warning"

            # 限制输出行数
            result["output"] = result["output"][-5:]  # 只保留最后5行

        except Exception as e:
            print(f"错误解析日志文件 {log_file}: {e}")
            result["status"] = "error"
            result["errors"].append(f"解析错误: {str(e)}")

        return result

    def analyze_all_logs(self) -> pd.DataFrame:
        """分析所有日志文件"""
        print(f"正在分析目录: {self.log_dir}")

        # 查找所有.log文件
        log_files = glob.glob(f"{self.log_dir}/*.log")

        if not log_files:
            print(f"在 {self.log_dir} 目录中未找到.log文件")
            return pd.DataFrame()

        print(f"找到 {len(log_files)} 个日志文件")

        all_results = []

        # 分析每个日志文件
        for i, log_file in enumerate(log_files, 1):
            print(f"分析进度: {i}/{len(log_files)} - {os.path.basename(log_file)}")
            result = self.parse_log_file(log_file)
            all_results.append(result)

            # 更新统计
            if result["technique"]:
                tech = result["technique"]
                stats = self.technique_stats[tech]
                stats["total"] += 1

                if result["status"] == "passed":
                    stats["passed"] += 1
                elif result["status"] == "failed":
                    stats["failed"] += 1
                elif result["status"] == "error":
                    stats["error"] += 1

                if result["execution_time"]:
                    stats["execution_times"].append(result["execution_time"])

                if result["indicators"]:
                    stats["findings"].extend(result["indicators"])

                if result["commands"]:
                    stats["commands"].extend(result["commands"])

                if result["timestamp"]:
                    stats["timestamps"].append(result["timestamp"])

                if result["error_type"]:
                    stats["error_types"].append(result["error_type"])

                stats["files"].append(result["filename"])

        self.results = all_results
        return pd.DataFrame(all_results)

    def calculate_statistics(self) -> dict:
        """计算统计信息"""
        if not self.results:
            return {}

        total = len(self.results)
        passed = sum(1 for r in self.results if r["status"] == "passed")
        failed = sum(1 for r in self.results if r["status"] == "failed")
        errors = sum(1 for r in self.results if r["status"] == "error")
        warnings = sum(1 for r in self.results if r["status"] == "warning")

        execution_times = [r["execution_time"] for r in self.results if r["execution_time"]]
        avg_execution_time = sum(execution_times) / len(execution_times) if execution_times else 0
        min_time = min(execution_times) if execution_times else 0
        max_time = max(execution_times) if execution_times else 0
        median_time = statistics.median(execution_times) if execution_times else 0

        techniques_tested = len(self.technique_stats)

        # 计算战术分类统计
        tactic_stats = defaultdict(lambda: {"total": 0, "passed": 0, "failed": 0, "techniques": set()})
        for result in self.results:
            tactic = result.get("tactic", "未知")
            tactic_stats[tactic]["total"] += 1
            tactic_stats[tactic]["techniques"].add(result.get("technique", "未知"))
            if result["status"] == "passed":
                tactic_stats[tactic]["passed"] += 1
            elif result["status"] in ["failed", "error"]:
                tactic_stats[tactic]["failed"] += 1

        # 转换为可序列化的格式
        tactic_summary = {}
        for tactic, data in tactic_stats.items():
            tactic_summary[tactic] = {
                "total": data["total"],
                "passed": data["passed"],
                "failed": data["failed"],
                "technique_count": len(data["techniques"]),
                "pass_rate": (data["passed"] / data["total"] * 100) if data["total"] > 0 else 0
            }

        # 计算风险分布
        risk_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0, "unknown": 0}
        risk_status = {"critical": {"passed": 0, "failed": 0}, "high": {"passed": 0, "failed": 0},
                      "medium": {"passed": 0, "failed": 0}, "low": {"passed": 0, "failed": 0},
                      "unknown": {"passed": 0, "failed": 0}}

        for result in self.results:
            tech = result["technique"]
            risk_level = "unknown"
            for level, techs in self.risk_levels.items():
                if tech in techs:
                    risk_level = level
                    break
            risk_counts[risk_level] += 1
            if result["status"] == "passed":
                risk_status[risk_level]["passed"] += 1
            elif result["status"] in ["failed", "error"]:
                risk_status[risk_level]["failed"] += 1

        # 计算错误类型分布
        error_types = Counter()
        for result in self.results:
            if result["error_type"]:
                error_types[result["error_type"]] += 1

        # 计算时间趋势
        hourly_distribution = defaultdict(lambda: {"total": 0, "passed": 0, "failed": 0})
        for result in self.results:
            if result["timestamp"]:
                hour = result["timestamp"][:13]  # 提取到小时
                hourly_distribution[hour]["total"] += 1
                if result["status"] == "passed":
                    hourly_distribution[hour]["passed"] += 1
                elif result["status"] in ["failed", "error"]:
                    hourly_distribution[hour]["failed"] += 1

        return {
            "total_tests": total,
            "passed": passed,
            "failed": failed,
            "errors": errors,
            "warnings": warnings,
            "pass_rate": (passed / total * 100) if total > 0 else 0,
            "avg_execution_time": avg_execution_time,
            "min_execution_time": min_time,
            "max_execution_time": max_time,
            "median_execution_time": median_time,
            "techniques_tested": techniques_tested,
            "risk_distribution": risk_counts,
            "risk_status": risk_status,
            "tactic_stats": tactic_summary,
            "error_types": dict(error_types),
            "hourly_distribution": dict(hourly_distribution)
        }

    def get_detailed_technique_analysis(self) -> dict:
        """获取技术详细分析"""
        analysis = {}
        for tech, stats in self.technique_stats.items():
            total = stats["total"]
            passed = stats["passed"]
            failed = stats["failed"]
            error = stats["error"]

            execution_times = stats["execution_times"]
            avg_time = sum(execution_times) / len(execution_times) if execution_times else 0
            min_time = min(execution_times) if execution_times else 0
            max_time = max(execution_times) if execution_times else 0

            # 计算稳定性 (变异系数)
            stability = "稳定"
            if len(execution_times) > 1:
                cv = statistics.stdev(execution_times) / avg_time if avg_time > 0 else 0
                if cv > 0.5:
                    stability = "不稳定"
                elif cv > 0.3:
                    stability = "一般"

            # 确定风险等级
            risk_level = "unknown"
            for level, techs in self.risk_levels.items():
                if tech in techs:
                    risk_level = level
                    break

            # 错误类型统计
            error_type_counts = Counter(stats["error_types"])

            analysis[tech] = {
                "description": self.technique_descriptions.get(tech, "未知技术"),
                "tactic": self.tactic_mapping.get(tech, "未知"),
                "total_tests": total,
                "passed": passed,
                "failed": failed,
                "error": error,
                "pass_rate": (passed / total * 100) if total > 0 else 0,
                "avg_execution_time": avg_time,
                "min_execution_time": min_time,
                "max_execution_time": max_time,
                "stability": stability,
                "risk_level": risk_level,
                "findings_count": len(stats["findings"]),
                "commands_count": len(stats["commands"]),
                "error_types": dict(error_type_counts),
                "files": stats["files"]
            }
        return analysis

    def get_attack_heatmap_data(self) -> dict:
        """生成ATT&CK覆盖热力图数据"""
        # 定义战术顺序（按ATT&CK矩阵顺序）
        tactic_order = [
            "初始访问", "执行", "持久化", "权限提升", "防御规避",
            "凭据访问", "发现", "横向移动", "数据收集", "数据外泄", "命令与控制"
        ]

        # 构建战术-技术矩阵
        heatmap_data = {}
        technique_by_tactic = defaultdict(list)

        for tech, stats in self.technique_stats.items():
            tactic = self.tactic_mapping.get(tech, "未知")
            total = stats["total"]
            passed = stats["passed"]
            pass_rate = (passed / total * 100) if total > 0 else 0

            technique_by_tactic[tactic].append({
                "technique": tech,
                "total": total,
                "passed": passed,
                "failed": stats["failed"],
                "error": stats["error"],
                "pass_rate": pass_rate,
                "description": self.technique_descriptions.get(tech, "未知技术")
            })

        # 按战术顺序组织数据
        ordered_tactics = []
        for tactic in tactic_order:
            if tactic in technique_by_tactic:
                ordered_tactics.append(tactic)

        # 添加未在预定义顺序中的战术
        for tactic in technique_by_tactic:
            if tactic not in ordered_tactics and tactic != "未知":
                ordered_tactics.append(tactic)
        if "未知" in technique_by_tactic:
            ordered_tactics.append("未知")

        # 构建热力图数据
        all_techniques = []
        heatmap_matrix = {}

        for tactic in ordered_tactics:
            techniques = sorted(technique_by_tactic[tactic], key=lambda x: x["technique"])
            heatmap_matrix[tactic] = {}
            for tech_data in techniques:
                tech = tech_data["technique"]
                if tech not in all_techniques:
                    all_techniques.append(tech)
                heatmap_matrix[tactic][tech] = {
                    "total": tech_data["total"],
                    "passed": tech_data["passed"],
                    "failed": tech_data["failed"],
                    "error": tech_data["error"],
                    "pass_rate": tech_data["pass_rate"],
                    "description": tech_data["description"]
                }

        # 计算战术汇总
        tactic_summary = {}
        for tactic in ordered_tactics:
            techs = technique_by_tactic[tactic]
            total = sum(t["total"] for t in techs)
            passed = sum(t["passed"] for t in techs)
            failed = sum(t["failed"] for t in techs)
            tactic_summary[tactic] = {
                "technique_count": len(techs),
                "total_tests": total,
                "passed": passed,
                "failed": failed,
                "pass_rate": (passed / total * 100) if total > 0 else 0
            }

        return {
            "tactics": ordered_tactics,
            "techniques": all_techniques,
            "matrix": heatmap_matrix,
            "technique_by_tactic": dict(technique_by_tactic),
            "tactic_summary": tactic_summary
        }

    def get_remediation_plan(self, stats: dict) -> dict:
        """生成详细的修复计划，包含具体实施步骤"""
        plan = {
            "immediate": [],
            "short_term": [],
            "medium_term": [],
            "long_term": []
        }

        # 立即行动项目
        critical_failed = stats["risk_status"]["critical"]["failed"]
        high_failed = stats["risk_status"]["high"]["failed"]

        # 获取失败的详细技术信息
        critical_failures = [r for r in self.results
                            if r["technique"] in self.risk_levels["critical"]
                            and r["status"] in ["failed", "error"]]
        high_failures = [r for r in self.results
                        if r["technique"] in self.risk_levels["high"]
                        and r["status"] in ["failed", "error"]]

        if critical_failed > 0:
            failed_techniques = list(set([r["technique"] for r in critical_failures]))[:5]
            plan["immediate"].append({
                "priority": "P0",
                "title": f"修复 {critical_failed} 个严重风险测试失败",
                "description": "严重风险测试失败表示关键安全控制失效，可能导致系统完全暴露于攻击之下",
                "impact": "如不及时修复，可能导致系统被完全攻破，数据泄露或服务中断",
                "action_steps": [
                    "立即识别失败的根本原因（配置错误/缺失安全控制/检测规则失效）",
                    "暂停相关系统的非必要访问，限制暴露面",
                    "部署临时防护措施（如额外防火墙规则、访问控制）",
                    "修复失败的安全控制配置",
                    "重新运行测试验证修复效果"
                ],
                "owners": ["安全架构师", "系统管理员", "安全运营团队"],
                "verification": "重新运行失败的测试用例，确保100%通过",
                "techniques": failed_techniques,
                "estimated_effort": "4-8小时"
            })

        if high_failed > 0:
            failed_techniques = list(set([r["technique"] for r in high_failures]))[:5]
            plan["immediate"].append({
                "priority": "P1",
                "title": f"审查 {high_failed} 个高风险测试失败",
                "description": "高风险测试失败可能严重影响系统安全，需优先处理和修复",
                "impact": "可能导致严重的安全漏洞，攻击者可利用这些漏洞获取系统控制权",
                "action_steps": [
                    "分析失败测试的技术类型和攻击向量",
                    "评估现有安全控制的有效性",
                    "实施针对性的安全加固措施",
                    "更新检测规则覆盖相关攻击技术",
                    "验证修复效果并记录经验"
                ],
                "owners": ["安全工程师", "应用管理员"],
                "verification": "高风险测试通过率提升至90%以上",
                "techniques": failed_techniques,
                "estimated_effort": "8-16小时"
            })

        # 按错误类型分析并生成修复建议
        error_types = stats["error_types"]
        if error_types:
            top_error = max(error_types.items(), key=lambda x: x[1])
            error_solutions = {
                "timeout": {
                    "steps": ["调整测试超时设置", "优化被测系统性能", "检查网络连接稳定性"],
                    "prevention": "建立性能基准监控，设置合理的超时阈值"
                },
                "connection_refused": {
                    "steps": ["验证服务运行状态", "检查防火墙规则", "确认端口配置正确"],
                    "prevention": "实施服务健康检查，建立自动重启机制"
                },
                "permission_denied": {
                    "steps": ["检查执行用户权限", "验证文件/目录权限设置", "确认SELinux/AppArmor策略"],
                    "prevention": "建立权限管理基线，定期审计权限配置"
                },
                "command_not_found": {
                    "steps": ["安装缺失依赖工具", "验证PATH环境变量", "确认工具版本兼容性"],
                    "prevention": "建立标准化环境，使用容器化部署测试"
                },
                "default": {
                    "steps": ["查看详细错误日志", "分析错误模式", "搜索相关解决方案"],
                    "prevention": "建立错误知识库，完善错误处理机制"
                }
            }
            solution = error_solutions.get(top_error[0], error_solutions["default"])

            plan["immediate"].append({
                "priority": "P1",
                "title": f"解决最常见的错误类型: {top_error[0]}",
                "description": f"共有 {top_error[1]} 个测试遇到此错误，需系统性解决以提升测试成功率",
                "impact": f"影响 {(top_error[1] / stats['total_tests'] * 100):.1f}% 的测试执行，降低测试覆盖率",
                "action_steps": solution["steps"],
                "owners": ["运维工程师", "测试工程师"],
                "verification": f"该错误出现次数降低80%以上",
                "prevention": solution["prevention"],
                "count": top_error[1],
                "estimated_effort": "4-12小时"
            })

        # 分析低通过率的战术类别
        low_pass_tactics = [(t, d) for t, d in stats['tactic_stats'].items()
                           if d['pass_rate'] < 50 and d['total'] >= 3]

        if low_pass_tactics:
            tactic_names = ', '.join([t[0] for t in low_pass_tactics[:3]])
            plan["short_term"].append({
                "priority": "P2",
                "title": f"强化低通过率战术类别: {tactic_names}",
                "description": f"共 {len(low_pass_tactics)} 个战术类别通过率低于50%，需针对性增强安全控制",
                "impact": "这些战术类别的安全覆盖不足，存在显著的检测和防御盲区",
                "action_steps": [
                    "分析低通过率战术的攻击场景",
                    "评估现有检测工具的配置和覆盖范围",
                    "部署或更新相关检测规则",
                    "配置自动响应和阻断机制",
                    "建立专项监控仪表板"
                ],
                "owners": ["安全运营中心(SOC)", "威胁情报团队"],
                "verification": "低通过率战术的通过率达到70%以上",
                "tactics": [t[0] for t in low_pass_tactics],
                "estimated_effort": "1-2周"
            })

        # 短期改进
        plan["short_term"].extend([
            {
                "priority": "P2",
                "title": "优化检测规则与告警策略",
                "description": f"基于 {stats['techniques_tested']} 种测试技术的检测结果，优化安全检测规则和告警策略",
                "impact": "减少误报和漏报，提高安全事件检测的准确性和时效性",
                "action_steps": [
                    "审查现有检测规则的覆盖范围",
                    "更新Sigma/YARA等检测规则",
                    "优化告警阈值，减少告警疲劳",
                    "建立测试用例与检测规则的映射关系",
                    "实施检测规则版本控制"
                ],
                "owners": ["安全工程师", "SOC分析师"],
                "verification": "检测规则覆盖率提升至90%以上，误报率降低50%",
                "estimated_effort": "2-3周"
            },
            {
                "priority": "P2",
                "title": "建立安全测试自动化流程",
                "description": "将MITRE ATT&CK测试集成到CI/CD流程，实现持续安全验证",
                "impact": "实现安全控制的持续验证，及时发现回归问题",
                "action_steps": [
                    "选择合适的自动化测试工具",
                    "编写自动化测试脚本和流水线",
                    "配置测试结果通知和报告机制",
                    "建立测试环境隔离和恢复机制",
                    "集成测试结果到安全仪表板"
                ],
                "owners": ["DevSecOps工程师", "安全架构师"],
                "verification": "自动化测试每周至少运行一次，覆盖率保持95%以上",
                "estimated_effort": "2-4周"
            },
            {
                "priority": "P2",
                "title": "开展针对性安全培训",
                "description": "基于测试结果中暴露的薄弱环节，对安全团队进行针对性培训",
                "impact": "提升团队对MITRE ATT&CK框架的理解和实际防御能力",
                "action_steps": [
                    "分析测试失败反映的知识缺口",
                    "设计针对性培训课程（技术分类）",
                    "组织红蓝对抗演练",
                    "建立内部知识库和案例集",
                    "评估培训效果并持续改进"
                ],
                "owners": ["安全培训负责人", "威胁狩猎团队"],
                "verification": "团队成员对相关战术类别的理解测试通过率达到80%",
                "estimated_effort": "1-2周"
            }
        ])

        # 中期改进
        plan["medium_term"].extend([
            {
                "priority": "P3",
                "title": "建立持续安全验证平台",
                "description": "建设统一的MITRE ATT&CK持续验证平台，支持多环境、多技术栈测试",
                "impact": "实现安全控制的标准化、常态化验证，提升整体安全成熟度",
                "action_steps": [
                    "评估和选型开源/商业测试平台",
                    "建设测试基础设施和隔离环境",
                    "开发定制化测试用例和场景",
                    "集成测试平台与SIEM/SOAR系统",
                    "建立测试结果度量和报告体系"
                ],
                "owners": ["安全架构师", "基础设施团队"],
                "verification": "平台支持全部ATT&CK战术类别测试，自动化率超过80%",
                "estimated_effort": "1-3个月"
            },
            {
                "priority": "P3",
                "title": "完善威胁检测与响应流程",
                "description": "基于测试结果优化安全事件检测、分析、响应全流程",
                "impact": "缩短威胁发现到响应的时间，提升事件处理效率",
                "action_steps": [
                    "绘制基于MITRE ATT&CK的检测覆盖矩阵",
                    "识别检测盲区并制定弥补方案",
                    "制定分类分级响应流程（SOP）",
                    "开展应急响应演练和桌面推演",
                    "建立威胁狩猎定期工作机制"
                ],
                "owners": ["安全运营经理", "事件响应团队"],
                "verification": "检测覆盖率达到80%以上，MTTR（平均响应时间）降低30%",
                "estimated_effort": "1-2个月"
            },
            {
                "priority": "P3",
                "title": "实施零信任安全架构",
                "description": "基于测试发现的横向移动和权限提升漏洞，推进零信任架构落地",
                "impact": "从根本上降低内部威胁风险，提升整体安全防御能力",
                "action_steps": [
                    "评估当前网络架构和访问控制现状",
                    "制定零信任实施路线图",
                    "部署微隔离和细粒度访问控制",
                    "实施持续身份验证和设备健康检查",
                    "建立动态访问策略引擎"
                ],
                "owners": ["网络架构师", "身份安全团队"],
                "verification": "关键系统实施微隔离，横向移动攻击测试100%被阻断",
                "estimated_effort": "2-3个月"
            }
        ])

        # 长期改进
        plan["long_term"].extend([
            {
                "priority": "P4",
                "title": "建立常态化红蓝对抗机制",
                "description": "建立定期的红蓝对抗演练机制，模拟真实APT攻击场景",
                "impact": "持续检验和提升整体安全防护能力，培养实战型安全人才",
                "action_steps": [
                    "制定红蓝对抗年度计划",
                    "建立红队（攻击方）和蓝队（防御方）",
                    "设计多阶段、多向量攻击场景",
                    "实施演练并进行全程记录",
                    "复盘总结并制定改进措施"
                ],
                "owners": ["安全总监", "红队/蓝队负责人"],
                "verification": "每季度至少开展1次红蓝对抗，平均检测时间逐步缩短",
                "estimated_effort": "持续进行"
            },
            {
                "priority": "P4",
                "title": "构建自适应安全架构",
                "description": "基于长期测试数据和威胁情报，构建智能化的自适应安全架构",
                "impact": "实现安全防护的自动化和智能化，主动应对新兴威胁",
                "action_steps": [
                    "部署UEBA（用户实体行为分析）系统",
                    "集成威胁情报自动更新检测规则",
                    "实施SOAR自动化响应编排",
                    "建立安全控制动态调整机制",
                    "开展AI/ML在安全检测中的应用"
                ],
                "owners": ["首席信息安全官", "安全研发团队"],
                "verification": "自动化响应率达到60%以上，新型威胁检测准确率达到85%",
                "estimated_effort": "3-6个月"
            },
            {
                "priority": "P4",
                "title": "参与行业安全生态共建",
                "description": "参与行业安全社区，贡献检测规则和最佳实践，提升行业整体安全水平",
                "impact": "提升组织安全影响力，获取前沿威胁情报和防御技术",
                "action_steps": [
                    "加入MITRE ATT&CK贡献者社区",
                    "开源内部开发的检测规则",
                    "参加行业安全会议和交流",
                    "与安全厂商建立合作",
                    "建立安全研究创新激励机制"
                ],
                "owners": ["安全战略团队", "外部合作部门"],
                "verification": "每年至少发布2个开源项目或技术分享",
                "estimated_effort": "持续进行"
            }
        ])

        return plan

    def generate_html_report(self, output_file: str = None):
        """生成HTML格式的完整报告"""
        if not output_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = f"{self.report_dir}/security_validation_report_windows_{timestamp}.html"

        stats = self.calculate_statistics()
        detailed_analysis = self.get_detailed_technique_analysis()
        remediation_plan = self.get_remediation_plan(stats)
        heatmap_data = self.get_attack_heatmap_data()

        # 创建HTML报告
        html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>主机安全验证测试报告 - {datetime.now().strftime('%Y-%m-%d')}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; line-height: 1.6; background-color: #f5f5f5; }}
        .header {{ background: linear-gradient(135deg, #2c3e50, #3498db); color: white; padding: 30px; border-radius: 10px; margin-bottom: 20px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }}
        .header h1 {{ margin: 0; font-size: 28px; }}
        .summary {{ background: white; padding: 25px; border-radius: 10px; margin-bottom: 20px; box-shadow: 0 2px 5px rgba(0,0,0,0.1); }}
        .section {{ background: white; padding: 25px; border-radius: 10px; margin-bottom: 20px; box-shadow: 0 2px 5px rgba(0,0,0,0.1); }}
        .section h2 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; margin-bottom: 20px; }}
        .section h3 {{ color: #34495e; margin-top: 20px; }}
        table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
        th {{ background-color: #34495e; color: white; font-weight: bold; }}
        tr:nth-child(even) {{ background-color: #f9f9f9; }}
        tr:hover {{ background-color: #f1f1f1; }}
        .passed {{ color: #27ae60; font-weight: bold; }}
        .failed {{ color: #e74c3c; font-weight: bold; }}
        .error {{ color: #f39c12; font-weight: bold; }}
        .warning {{ color: #ff9800; font-weight: bold; }}
        .critical {{ background-color: #ffebee; color: #c62828; font-weight: bold; }}
        .high {{ background-color: #fff3e0; color: #ef6c00; font-weight: bold; }}
        .medium {{ background-color: #fffde7; color: #f9a825; font-weight: bold; }}
        .low {{ background-color: #e8f5e9; color: #2e7d32; font-weight: bold; }}
        .unknown {{ background-color: #f5f5f5; color: #616161; }}
        .chart-container {{ margin: 25px 0; padding: 20px; background: #fafafa; border-radius: 8px; }}
        .progress-bar {{ width: 100%; background-color: #e0e0e0; border-radius: 10px; overflow: hidden; margin: 10px 0; height: 30px; }}
        .progress-fill {{ height: 100%; background: linear-gradient(90deg, #4CAF50, #8BC34A); text-align: center; line-height: 30px; color: white; font-weight: bold; }}
        .details {{ display: none; margin-top: 15px; padding: 20px; background: #f8f9fa; border-left: 4px solid #3498db; border-radius: 5px; }}
        .toggle-btn {{ background: #3498db; color: white; border: none; padding: 8px 16px; cursor: pointer; margin: 5px 0; border-radius: 5px; transition: background 0.3s; }}
        .toggle-btn:hover {{ background: #2980b9; }}
        .metric-box {{ display: inline-block; width: 18%; margin: 1%; padding: 20px; background: #ecf0f1; border-radius: 10px; text-align: center; }}
        .metric-box h4 {{ margin: 0; color: #7f8c8d; font-size: 14px; }}
        .metric-box .value {{ font-size: 32px; font-weight: bold; color: #2c3e50; margin: 10px 0; }}
        .remediation-item {{ padding: 15px; margin: 10px 0; border-radius: 8px; border-left: 4px solid; }}
        .remediation-immediate {{ background: #ffebee; border-color: #e74c3c; }}
        .remediation-short {{ background: #fff3e0; border-color: #f39c12; }}
        .remediation-medium {{ background: #e3f2fd; border-color: #3498db; }}
        .remediation-long {{ background: #e8f5e9; border-color: #27ae60; }}
        .priority-badge {{ display: inline-block; padding: 4px 12px; border-radius: 12px; font-size: 12px; font-weight: bold; margin-right: 10px; }}
        .priority-p0 {{ background: #e74c3c; color: white; }}
        .priority-p1 {{ background: #f39c12; color: white; }}
        .priority-p2 {{ background: #3498db; color: white; }}
        .priority-p3 {{ background: #9b59b6; color: white; }}
        .priority-p4 {{ background: #95a5a6; color: white; }}
        .timeline {{ position: relative; padding-left: 30px; }}
        .timeline::before {{ content: ''; position: absolute; left: 10px; top: 0; bottom: 0; width: 2px; background: #3498db; }}
        .timeline-item {{ position: relative; margin: 15px 0; padding: 15px; background: #f8f9fa; border-radius: 8px; }}
        .timeline-item::before {{ content: ''; position: absolute; left: -24px; top: 20px; width: 12px; height: 12px; border-radius: 50%; background: #3498db; }}
        .stats-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 20px; margin: 20px 0; }}
        .stat-card {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 10px; }}
        .stat-card h4 {{ margin: 0 0 10px 0; opacity: 0.9; }}
        .stat-card .number {{ font-size: 36px; font-weight: bold; }}
        .heatmap-container {{ margin: 20px 0; overflow-x: auto; }}
        .heatmap-table {{ border-collapse: collapse; font-size: 12px; min-width: 100%; }}
        .heatmap-table th, .heatmap-table td {{ border: 1px solid #ddd; padding: 8px; text-align: center; min-width: 60px; }}
        .heatmap-table th {{ background-color: #34495e; color: white; font-weight: bold; position: sticky; left: 0; z-index: 10; }}
        .heatmap-table .tactic-header {{ background-color: #2c3e50; text-align: left; min-width: 120px; position: sticky; left: 0; z-index: 20; }}
        .heatmap-table .tactic-summary {{ background-color: #ecf0f1; font-weight: bold; }}
        .heatmap-cell {{ cursor: pointer; transition: transform 0.1s; }}
        .heatmap-cell:hover {{ transform: scale(1.1); box-shadow: 0 2px 8px rgba(0,0,0,0.3); z-index: 5; position: relative; }}
        .heatmap-legend {{ display: flex; align-items: center; margin: 15px 0; padding: 10px; background: #f8f9fa; border-radius: 5px; }}
        .heatmap-legend-item {{ display: flex; align-items: center; margin-right: 20px; }}
        .heatmap-legend-color {{ width: 20px; height: 20px; margin-right: 5px; border-radius: 3px; }}
        .rate-excellent {{ background-color: #27ae60; color: white; }}
        .rate-good {{ background-color: #58d68d; color: white; }}
        .rate-medium {{ background-color: #f39c12; color: white; }}
        .rate-poor {{ background-color: #e74c3c; color: white; }}
        .rate-critical {{ background-color: #c0392b; color: white; }}
        .rate-notested {{ background-color: #bdc3c7; color: #666; }}
        .heatmap-tooltip {{ position: absolute; background: rgba(0,0,0,0.85); color: white; padding: 10px; border-radius: 5px; font-size: 11px; z-index: 100; max-width: 250px; display: none; }}
    </style>
    <script src="https://cdn.jsdelivr.net/npm/chart.js"></script>
</head>
<body>
    <div class="header">
        <h1>MITRE ATT&CK 主机安全验证测试报告</h1>
        <p>报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p>测试环境: Windows 系统 | 工具: goart-windows</p>
        <p>测试范围: {stats['techniques_tested']} 种技术 | {stats['total_tests']} 次测试</p>
    </div>

    <div class="summary">
        <h2>执行摘要</h2>
        <div class="stats-grid">
            <div class="stat-card">
                <h4>总测试数</h4>
                <div class="number">{stats['total_tests']}</div>
            </div>
            <div class="stat-card" style="background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%);">
                <h4>通过测试</h4>
                <div class="number">{stats['passed']}</div>
            </div>
            <div class="stat-card" style="background: linear-gradient(135deg, #eb3349 0%, #f45c43 100%);">
                <h4>失败测试</h4>
                <div class="number">{stats['failed']}</div>
            </div>
            <div class="stat-card" style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);">
                <h4>错误测试</h4>
                <div class="number">{stats['errors']}</div>
            </div>
        </div>

        <h3>通过率分析</h3>
        <div class="progress-bar">
            <div class="progress-fill" style="width: {stats['pass_rate']}%;">{stats['pass_rate']:.1f}%</div>
        </div>
        <p>整体通过率: <strong>{stats['pass_rate']:.1f}%</strong> |
           平均执行时间: <strong>{stats['avg_execution_time']:.2f}秒</strong> |
           最短: <strong>{stats['min_execution_time']:.2f}秒</strong> |
           最长: <strong>{stats['max_execution_time']:.2f}秒</strong> |
           中位数: <strong>{stats['median_execution_time']:.2f}秒</strong>
        </p>
    </div>

    <div class="section">
        <h2>风险评估矩阵</h2>
        <table>
            <tr>
                <th>风险等级</th>
                <th>测试数量</th>
                <th>通过数</th>
                <th>失败数</th>
                <th>通过率</th>
                <th>占比</th>
                <th>状态评估</th>
            </tr>
            <tr class="critical">
                <td>严重</td>
                <td>{stats['risk_distribution']['critical']}</td>
                <td class="passed">{stats['risk_status']['critical']['passed']}</td>
                <td class="failed">{stats['risk_status']['critical']['failed']}</td>
                <td>{(stats['risk_status']['critical']['passed']/stats['risk_distribution']['critical']*100) if stats['risk_distribution']['critical'] > 0 else 0:.1f}%</td>
                <td>{(stats['risk_distribution']['critical']/stats['total_tests']*100) if stats['total_tests'] > 0 else 0:.1f}%</td>
                <td>{"✓ 良好" if stats['risk_status']['critical']['failed'] == 0 else "⚠ 需立即关注"}</td>
            </tr>
            <tr class="high">
                <td>高</td>
                <td>{stats['risk_distribution']['high']}</td>
                <td class="passed">{stats['risk_status']['high']['passed']}</td>
                <td class="failed">{stats['risk_status']['high']['failed']}</td>
                <td>{(stats['risk_status']['high']['passed']/stats['risk_distribution']['high']*100) if stats['risk_distribution']['high'] > 0 else 0:.1f}%</td>
                <td>{(stats['risk_distribution']['high']/stats['total_tests']*100) if stats['total_tests'] > 0 else 0:.1f}%</td>
                <td>{"✓ 可控" if stats['risk_status']['high']['failed'] == 0 else "⚠ 需优先处理"}</td>
            </tr>
            <tr class="medium">
                <td>中</td>
                <td>{stats['risk_distribution']['medium']}</td>
                <td class="passed">{stats['risk_status']['medium']['passed']}</td>
                <td class="failed">{stats['risk_status']['medium']['failed']}</td>
                <td>{(stats['risk_status']['medium']['passed']/stats['risk_distribution']['medium']*100) if stats['risk_distribution']['medium'] > 0 else 0:.1f}%</td>
                <td>{(stats['risk_distribution']['medium']/stats['total_tests']*100) if stats['total_tests'] > 0 else 0:.1f}%</td>
                <td>正常监控</td>
            </tr>
            <tr class="low">
                <td>低</td>
                <td>{stats['risk_distribution']['low']}</td>
                <td class="passed">{stats['risk_status']['low']['passed']}</td>
                <td class="failed">{stats['risk_status']['low']['failed']}</td>
                <td>{(stats['risk_status']['low']['passed']/stats['risk_distribution']['low']*100) if stats['risk_distribution']['low'] > 0 else 0:.1f}%</td>
                <td>{(stats['risk_distribution']['low']/stats['total_tests']*100) if stats['total_tests'] > 0 else 0:.1f}%</td>
                <td>建议关注</td>
            </tr>
        </table>
    </div>

    <div class="section">
        <h2>战术分类统计</h2>
        <table>
            <tr>
                <th>战术类别</th>
                <th>测试次数</th>
                <th>技术数量</th>
                <th>通过数</th>
                <th>失败数</th>
                <th>通过率</th>
                <th>安全状态</th>
            </tr>
"""

        # 添加战术统计行
        for tactic, tactic_data in sorted(stats['tactic_stats'].items(), key=lambda x: x[1]['total'], reverse=True):
            security_status = "✓ 良好" if tactic_data['pass_rate'] >= 80 else ("⚠ 一般" if tactic_data['pass_rate'] >= 50 else "✗ 需改进")
            html_content += f"""
            <tr>
                <td><strong>{tactic}</strong></td>
                <td>{tactic_data['total']}</td>
                <td>{tactic_data['technique_count']}</td>
                <td class="passed">{tactic_data['passed']}</td>
                <td class="failed">{tactic_data['failed']}</td>
                <td>{tactic_data['pass_rate']:.1f}%</td>
                <td>{security_status}</td>
            </tr>
"""

        html_content += """
        </table>
    </div>
"""

        # 添加错误类型分析（如果有错误）
        if stats['error_types']:
            html_content += f"""
    <div class="section">
        <h2>错误类型分析</h2>
        <table>
            <tr>
                <th>错误类型</th>
                <th>出现次数</th>
                <th>占比</th>
                <th>建议处理措施</th>
            </tr>
"""
            total_errors = sum(stats['error_types'].values())
            for error_type, count in sorted(stats['error_types'].items(), key=lambda x: x[1], reverse=True):
                percentage = (count / total_errors * 100) if total_errors > 0 else 0
                recommendations = {
                    "权限不足": "检查测试账户权限，确保具有执行所需权限",
                    "命令未找到": "安装缺失的依赖工具或更新PATH环境变量",
                    "网络错误": "检查网络连接和防火墙规则配置",
                    "资源不足": "增加系统资源或优化测试执行计划",
                    "依赖缺失": "安装所需的Python模块或系统库",
                    "配置错误": "检查配置文件和参数设置",
                    "语法错误": "检查测试脚本语法和格式"
                }
                html_content += f"""
            <tr>
                <td>{error_type}</td>
                <td>{count}</td>
                <td>{percentage:.1f}%</td>
                <td>{recommendations.get(error_type, "请检查相关配置和日志")}</td>
            </tr>
"""
            html_content += """
        </table>
    </div>
"""

        # 添加详细技术分析
        html_content += """
    <div class="section">
        <h2>详细技术分析</h2>
        <table>
            <tr>
                <th>技术ID</th>
                <th>战术类别</th>
                <th>描述</th>
                <th>测试次数</th>
                <th>通过率</th>
                <th>平均时间(秒)</th>
                <th>时间范围</th>
                <th>稳定性</th>
                <th>风险等级</th>
                <th>发现指标</th>
                <th>操作</th>
            </tr>
"""

        for tech, analysis in sorted(detailed_analysis.items(), key=lambda x: x[1]['total_tests'], reverse=True):
            status_class = "passed" if analysis['pass_rate'] >= 80 else ("warning" if analysis['pass_rate'] >= 50 else "failed")
            risk_class = analysis['risk_level']

            html_content += f"""
            <tr>
                <td><strong>{tech}</strong></td>
                <td>{analysis['tactic']}</td>
                <td>{analysis['description'][:50]}...</td>
                <td>{analysis['total_tests']}</td>
                <td class="{status_class}">{analysis['pass_rate']:.1f}%</td>
                <td>{analysis['avg_execution_time']:.2f}</td>
                <td>{analysis['min_execution_time']:.2f} - {analysis['max_execution_time']:.2f}</td>
                <td>{analysis['stability']}</td>
                <td class="{risk_class}">{risk_class}</td>
                <td>{analysis['findings_count']}</td>
                <td><button class="toggle-btn" onclick="toggleDetails('tech_{tech}')">查看详情</button></td>
            </tr>
            <tr id="details_tech_{tech}" class="details-row">
                <td colspan="11">
                    <div id="detail_tech_{tech}" class="details">
                        <h4>{tech} - {analysis['description']}</h4>
                        <p><strong>战术类别:</strong> {analysis['tactic']}</p>
                        <p><strong>测试文件:</strong> {', '.join(analysis['files'][:3])}</p>
                        <p><strong>执行命令数:</strong> {analysis['commands_count']}</p>
                        <p><strong>错误类型分布:</strong> {analysis['error_types']}</p>
                        <p><strong>性能分析:</strong> 平均 {analysis['avg_execution_time']:.2f}秒,
                           最短 {analysis['min_execution_time']:.2f}秒,
                           最长 {analysis['max_execution_time']:.2f}秒</p>
                    </div>
                </td>
            </tr>
"""

        html_content += """
        </table>
    </div>
"""

        # 添加修复建议
        html_content += """
    <div class="section">
        <h2>修复建议与行动计划</h2>

        <h3>立即行动（0-24小时）</h3>
"""
        for item in remediation_plan['immediate']:
            techniques = item.get('techniques', [])
            tech_info = f"涉及技术: {', '.join(techniques[:5])}" if techniques else ""
            html_content += f"""
        <div class="remediation-item remediation-immediate">
            <span class="priority-badge priority-{item['priority'].lower()}">{item['priority']}</span>
            <strong>{item['title']}</strong>
            <p>{item['description']}</p>
            <p style="color: #666; font-size: 12px;">{tech_info}</p>
        </div>
"""

        html_content += """
        <h3>短期改进（1-2周）</h3>
"""
        for item in remediation_plan['short_term']:
            html_content += f"""
        <div class="remediation-item remediation-short">
            <span class="priority-badge priority-{item['priority'].lower()}">{item['priority']}</span>
            <strong>{item['title']}</strong>
            <p>{item['description']}</p>
        </div>
"""

        html_content += """
        <h3>中期改进（1-3个月）</h3>
"""
        for item in remediation_plan['medium_term']:
            html_content += f"""
        <div class="remediation-item remediation-medium">
            <span class="priority-badge priority-{item['priority'].lower()}">{item['priority']}</span>
            <strong>{item['title']}</strong>
            <p>{item['description']}</p>
        </div>
"""

        html_content += """
        <h3>长期优化（3个月以上）</h3>
"""
        for item in remediation_plan['long_term']:
            html_content += f"""
        <div class="remediation-item remediation-long">
            <span class="priority-badge priority-{item['priority'].lower()}">{item['priority']}</span>
            <strong>{item['title']}</strong>
            <p>{item['description']}</p>
        </div>
"""

        html_content += """
    </div>

    <div class="section">
        <h2>测试执行图表</h2>
        <div class="chart-container">
            <canvas id="resultsChart" width="400" height="200"></canvas>
        </div>
        <div class="chart-container">
            <canvas id="riskChart" width="400" height="200"></canvas>
        </div>
        <div class="chart-container">
            <canvas id="tacticChart" width="400" height="200"></canvas>
        </div>
    </div>

    <div class="section">
        <h2>ATT&CK 覆盖热力图</h2>
        <p style="color: #666; margin-bottom: 15px;">热力图展示各战术下技术的测试覆盖情况和通过率。颜色越绿表示通过率越高，颜色越红表示通过率越低或存在风险。</p>

        <div class="heatmap-legend">
            <div class="heatmap-legend-item">
                <div class="heatmap-legend-color rate-excellent"></div>
                <span>≥90% 优秀</span>
            </div>
            <div class="heatmap-legend-item">
                <div class="heatmap-legend-color rate-good"></div>
                <span>70-89% 良好</span>
            </div>
            <div class="heatmap-legend-item">
                <div class="heatmap-legend-color rate-medium"></div>
                <span>50-69% 一般</span>
            </div>
            <div class="heatmap-legend-item">
                <div class="heatmap-legend-color rate-poor"></div>
                <span>25-49% 较差</span>
            </div>
            <div class="heatmap-legend-item">
                <div class="heatmap-legend-color rate-critical"></div>
                <span>&lt;25% 严重</span>
            </div>
            <div class="heatmap-legend-item">
                <div class="heatmap-legend-color rate-notested"></div>
                <span>未测试</span>
            </div>
        </div>

        <div class="heatmap-container">
            <table class="heatmap-table" id="attackHeatmap">
                <thead>
                    <tr>
                        <th class="tactic-header">战术类别</th>
                        <th>技术数</th>
                        <th>测试数</th>
                        <th>通过数</th>
                        <th>失败数</th>
                        <th>通过率</th>
                        <th colspan="10">技术覆盖详情（点击查看）</th>
                    </tr>
                </thead>
                <tbody>
"""

        # 生成热力图数据行
        for tactic in heatmap_data['tactics']:
            tactic_info = heatmap_data['tactic_summary'].get(tactic, {})
            tech_list = heatmap_data['technique_by_tactic'].get(tactic, [])

            # 获取通过率对应的样式类
            pass_rate = tactic_info.get('pass_rate', 0)
            if pass_rate >= 90:
                rate_class = "rate-excellent"
            elif pass_rate >= 70:
                rate_class = "rate-good"
            elif pass_rate >= 50:
                rate_class = "rate-medium"
            elif pass_rate >= 25:
                rate_class = "rate-poor"
            elif pass_rate > 0:
                rate_class = "rate-critical"
            else:
                rate_class = "rate-notested"

            html_content += f"""
                    <tr>
                        <td class="tactic-header">{tactic}</td>
                        <td class="tactic-summary">{tactic_info.get('technique_count', 0)}</td>
                        <td class="tactic-summary">{tactic_info.get('total_tests', 0)}</td>
                        <td class="tactic-summary">{tactic_info.get('passed', 0)}</td>
                        <td class="tactic-summary">{tactic_info.get('failed', 0)}</td>
                        <td class="tactic-summary {rate_class}">{pass_rate:.1f}%</td>
"""

            # 添加技术单元格（最多显示10个）
            displayed_techs = sorted(tech_list, key=lambda x: x['pass_rate'])[:10]
            for tech_data in displayed_techs:
                tech_pass_rate = tech_data['pass_rate']
                if tech_pass_rate >= 90:
                    tech_class = "rate-excellent"
                elif tech_pass_rate >= 70:
                    tech_class = "rate-good"
                elif tech_pass_rate >= 50:
                    tech_class = "rate-medium"
                elif tech_pass_rate >= 25:
                    tech_class = "rate-poor"
                else:
                    tech_class = "rate-critical"

                html_content += f"""
                        <td class="heatmap-cell {tech_class}" title="{tech_data['technique']}: {tech_data['description'][:50]}...&#10;通过率: {tech_pass_rate:.1f}%&#10;测试数: {tech_data['total']}">
                            {tech_data['technique']}<br><small>{tech_pass_rate:.0f}%</small>
                        </td>
"""

            # 填充空白单元格
            remaining_cells = 10 - len(displayed_techs)
            for _ in range(remaining_cells):
                html_content += """                        <td class="rate-notested">-</td>\n"""

            html_content += """                    </tr>\n"""

        html_content += """
                </tbody>
            </table>
        </div>

        <div style="margin-top: 15px; padding: 15px; background: #f8f9fa; border-radius: 5px;">
            <h4>热力图解读说明</h4>
            <ul style="margin: 10px 0; padding-left: 20px;">
                <li><strong>颜色编码：</strong>绿色表示测试通过率高（安全控制有效），红色表示通过率低（存在安全风险）</li>
                <li><strong>战术覆盖：</strong>每个战术行展示该类别下测试的技术数量和通过率</li>
                <li><strong>技术详情：</strong>鼠标悬停在技术单元格上可查看详细信息</li>
                <li><strong>优先处理：</strong>重点关注红色和橙色单元格对应的技术</li>
            </ul>
        </div>
    </div>

    <div class="section">
        <h2>测试结果原始数据</h2>
        <table>
            <tr>
                <th>技术ID</th>
                <th>技术描述</th>
                <th>战术类别</th>
                <th>状态</th>
                <th>执行时间(秒)</th>
                <th>检测指标</th>
                <th>错误类型</th>
                <th>风险等级</th>
            </tr>
"""

        for result in self.results:
            tech = result["technique"] or "未知"
            description = result["description"] or "未知技术"
            status = result["status"]
            tactic = result.get("tactic", "未知")
            exec_time = result["execution_time"] or "N/A"
            indicators_count = len(result["indicators"])
            error_type = result.get("error_type", "N/A")

            risk_level = "unknown"
            for level, techs in self.risk_levels.items():
                if tech in techs:
                    risk_level = level
                    break

            status_class = status

            html_content += f"""
            <tr>
                <td>{tech}</td>
                <td>{description}</td>
                <td>{tactic}</td>
                <td class="{status_class}">{status}</td>
                <td>{exec_time}</td>
                <td>{indicators_count}</td>
                <td>{error_type}</td>
                <td class="{risk_level}">{risk_level}</td>
            </tr>
"""

        html_content += f"""
        </table>
    </div>

    <script>
        // 测试结果分布图
        const ctx1 = document.getElementById('resultsChart').getContext('2d');
        new Chart(ctx1, {{
            type: 'doughnut',
            data: {{
                labels: ['通过', '失败', '错误', '警告'],
                datasets: [{{
                    data: [{stats['passed']}, {stats['failed']}, {stats['errors']}, {stats['warnings']}],
                    backgroundColor: ['#28a745', '#dc3545', '#ffc107', '#17a2b8'],
                    borderWidth: 2
                }}]
            }},
            options: {{
                responsive: true,
                plugins: {{
                    title: {{ display: true, text: '测试结果分布', font: {{ size: 18 }} }},
                    legend: {{ position: 'bottom' }}
                }}
            }}
        }});

        // 风险等级分布图
        const ctx2 = document.getElementById('riskChart').getContext('2d');
        new Chart(ctx2, {{
            type: 'bar',
            data: {{
                labels: ['严重', '高', '中', '低', '未知'],
                datasets: [
                    {{
                        label: '通过',
                        data: [{stats['risk_status']['critical']['passed']}, {stats['risk_status']['high']['passed']}, {stats['risk_status']['medium']['passed']}, {stats['risk_status']['low']['passed']}, 0],
                        backgroundColor: '#28a745'
                    }},
                    {{
                        label: '失败',
                        data: [{stats['risk_status']['critical']['failed']}, {stats['risk_status']['high']['failed']}, {stats['risk_status']['medium']['failed']}, {stats['risk_status']['low']['failed']}, 0],
                        backgroundColor: '#dc3545'
                    }}
                ]
            }},
            options: {{
                responsive: true,
                plugins: {{
                    title: {{ display: true, text: '风险等级测试结果', font: {{ size: 18 }} }}
                }},
                scales: {{
                    x: {{ stacked: true }},
                    y: {{ stacked: true }}
                }}
            }}
        }});

        // 战术分类通过率图
        const ctx3 = document.getElementById('tacticChart').getContext('2d');
        new Chart(ctx3, {{
            type: 'radar',
            data: {{
                labels: {list(stats['tactic_stats'].keys())},
                datasets: [{{
                    label: '通过率(%)',
                    data: {[t['pass_rate'] for t in stats['tactic_stats'].values()]},
                    backgroundColor: 'rgba(52, 152, 219, 0.2)',
                    borderColor: '#3498db',
                    pointBackgroundColor: '#3498db'
                }}]
            }},
            options: {{
                responsive: true,
                plugins: {{
                    title: {{ display: true, text: '战术分类安全评分', font: {{ size: 18 }} }}
                }},
                scales: {{
                    r: {{
                        beginAtZero: true,
                        max: 100
                    }}
                }}
            }}
        }});

        // 切换详情显示
        function toggleDetails(id) {{
            const detailDiv = document.getElementById('detail_' + id);
            if (detailDiv.style.display === 'block') {{
                detailDiv.style.display = 'none';
            }} else {{
                detailDiv.style.display = 'block';
            }}
        }}
    </script>
</body>
</html>
"""

        # 保存HTML文件
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html_content)

        print(f"HTML报告已生成: {output_file}")
        return output_file

    def generate_excel_report(self, output_file: str = None):
        """生成Excel格式的详细报告"""
        if not output_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = f"{self.report_dir}/security_validation_report_windows_{timestamp}.xlsx"

        # 创建详细结果DataFrame
        detailed_data = []
        for result in self.results:
            risk_level = "unknown"
            for level, techs in self.risk_levels.items():
                if result["technique"] in techs:
                    risk_level = level
                    break

            detailed_data.append({
                "技术ID": result["technique"] or "未知",
                "技术描述": result["description"] or "未知",
                "战术类别": result.get("tactic", "未知"),
                "文件名": result["filename"],
                "时间戳": result["timestamp"] or "",
                "状态": result["status"],
                "执行时间(秒)": result["execution_time"] or 0,
                "行数": result.get("line_count", 0),
                "输出大小(B)": result.get("output_size", 0),
                "执行命令数量": len(result["commands"]),
                "检测指标数量": len(result["indicators"]),
                "警告数量": len(result["warnings"]),
                "错误数量": len(result["errors"]),
                "错误类型": result.get("error_type", "N/A"),
                "风险等级": risk_level,
                "执行命令": "\n".join(result["commands"][:5]),
                "检测指标": "\n".join(result["indicators"][:5]),
                "警告信息": "\n".join(result["warnings"][:3]),
                "错误信息": "\n".join(result["errors"][:3])
            })

        detailed_df = pd.DataFrame(detailed_data)

        # 创建技术统计DataFrame
        detailed_analysis = self.get_detailed_technique_analysis()
        stats_data = []
        for tech, analysis in detailed_analysis.items():
            stats_data.append({
                "技术ID": tech,
                "技术描述": analysis["description"],
                "战术类别": analysis["tactic"],
                "测试次数": analysis["total_tests"],
                "通过次数": analysis["passed"],
                "失败次数": analysis["failed"],
                "错误次数": analysis["error"],
                "通过率(%)": analysis["pass_rate"],
                "平均执行时间(秒)": analysis["avg_execution_time"],
                "最短执行时间(秒)": analysis["min_execution_time"],
                "最长执行时间(秒)": analysis["max_execution_time"],
                "稳定性": analysis["stability"],
                "风险等级": analysis["risk_level"],
                "发现指标数量": analysis["findings_count"],
                "执行命令数量": analysis["commands_count"]
            })

        stats_df = pd.DataFrame(stats_data)

        # 创建汇总统计DataFrame
        summary_stats = self.calculate_statistics()
        summary_data = [{
            "总测试数": summary_stats["total_tests"],
            "通过测试": summary_stats["passed"],
            "失败测试": summary_stats["failed"],
            "错误测试": summary_stats["errors"],
            "警告测试": summary_stats["warnings"],
            "通过率(%)": summary_stats["pass_rate"],
            "平均执行时间(秒)": summary_stats["avg_execution_time"],
            "最短执行时间(秒)": summary_stats["min_execution_time"],
            "最长执行时间(秒)": summary_stats["max_execution_time"],
            "中位数执行时间(秒)": summary_stats["median_execution_time"],
            "测试技术数量": summary_stats["techniques_tested"],
            "严重风险测试": summary_stats["risk_distribution"]["critical"],
            "高风险测试": summary_stats["risk_distribution"]["high"],
            "中等风险测试": summary_stats["risk_distribution"]["medium"],
            "低风险测试": summary_stats["risk_distribution"]["low"],
            "严重风险通过": summary_stats["risk_status"]["critical"]["passed"],
            "严重风险失败": summary_stats["risk_status"]["critical"]["failed"]
        }]

        summary_df = pd.DataFrame(summary_data)

        # 创建战术统计DataFrame
        tactic_data = []
        for tactic, data in summary_stats["tactic_stats"].items():
            tactic_data.append({
                "战术类别": tactic,
                "测试次数": data["total"],
                "技术数量": data["technique_count"],
                "通过数": data["passed"],
                "失败数": data["failed"],
                "通过率(%)": data["pass_rate"]
            })

        tactic_df = pd.DataFrame(tactic_data)

        # 创建错误类型统计DataFrame
        error_data = []
        for error_type, count in summary_stats["error_types"].items():
            error_data.append({
                "错误类型": error_type,
                "出现次数": count,
                "占比(%)": (count / sum(summary_stats["error_types"].values()) * 100) if summary_stats["error_types"] else 0
            })

        error_df = pd.DataFrame(error_data)

        # 创建修复建议DataFrame
        remediation_plan = self.get_remediation_plan(summary_stats)
        recommendations_data = []
        for priority, items in [("P0-立即", remediation_plan["immediate"]),
                                ("P1-紧急", remediation_plan["immediate"][1:] if len(remediation_plan["immediate"]) > 1 else []),
                                ("P2-短期", remediation_plan["short_term"]),
                                ("P3-中期", remediation_plan["medium_term"]),
                                ("P4-长期", remediation_plan["long_term"])]:
            for item in items:
                recommendations_data.append({
                    "优先级": priority,
                    "标题": item["title"],
                    "描述": item["description"]
                })

        recommendations_df = pd.DataFrame(recommendations_data)

        # 创建ATT&CK热力图数据DataFrame
        heatmap_data = self.get_attack_heatmap_data()
        heatmap_rows = []
        for tactic in heatmap_data['tactics']:
            tactic_info = heatmap_data['tactic_summary'].get(tactic, {})
            tech_list = heatmap_data['technique_by_tactic'].get(tactic, [])

            # 战术汇总行
            row_data = {
                "战术类别": tactic,
                "技术数量": tactic_info.get('technique_count', 0),
                "测试总数": tactic_info.get('total_tests', 0),
                "通过数": tactic_info.get('passed', 0),
                "失败数": tactic_info.get('failed', 0),
                "通过率(%)": round(tactic_info.get('pass_rate', 0), 1)
            }

            # 添加技术详情列（最多10个）
            for i, tech_data in enumerate(sorted(tech_list, key=lambda x: x['pass_rate'])[:10], 1):
                row_data[f"技术{i}"] = tech_data['technique']
                row_data[f"技术{i}通过率(%)"] = round(tech_data['pass_rate'], 1)

            heatmap_rows.append(row_data)

        heatmap_df = pd.DataFrame(heatmap_rows)

        try:
            # 保存到Excel
            with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
                summary_df.to_excel(writer, sheet_name='执行摘要', index=False)
                detailed_df.to_excel(writer, sheet_name='详细结果', index=False)
                stats_df.to_excel(writer, sheet_name='技术统计', index=False)
                tactic_df.to_excel(writer, sheet_name='战术统计', index=False)
                error_df.to_excel(writer, sheet_name='错误分析', index=False)
                recommendations_df.to_excel(writer, sheet_name='修复建议', index=False)
                heatmap_df.to_excel(writer, sheet_name='ATT&CK热力图', index=False)

            print(f"Excel报告已生成: {output_file}")
            return output_file

        except Exception as e:
            print(f"生成Excel报告时出错: {e}")
            import traceback
            traceback.print_exc()
            return None

    def generate_text_report(self, output_file: str = None):
        """生成文本格式的详细报告"""
        if not output_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = f"{self.report_dir}/security_validation_report_windows_{timestamp}.txt"

        stats = self.calculate_statistics()
        detailed_analysis = self.get_detailed_technique_analysis()
        remediation_plan = self.get_remediation_plan(stats)

        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("=" * 100 + "\n")
            f.write("MITRE ATT&CK 安全验证测试详细报告\n")
            f.write("=" * 100 + "\n\n")

            f.write("【执行摘要】\n")
            f.write("-" * 80 + "\n")
            f.write(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"总测试数: {stats['total_tests']}\n")
            f.write(f"通过测试: {stats['passed']} ({stats['pass_rate']:.1f}%)\n")
            f.write(f"失败测试: {stats['failed']}\n")
            f.write(f"错误测试: {stats['errors']}\n")
            f.write(f"警告测试: {stats['warnings']}\n")
            f.write(f"测试技术数量: {stats['techniques_tested']}种\n\n")

            f.write("【执行时间统计】\n")
            f.write("-" * 80 + "\n")
            f.write(f"平均执行时间: {stats['avg_execution_time']:.2f}秒\n")
            f.write(f"最短执行时间: {stats['min_execution_time']:.2f}秒\n")
            f.write(f"最长执行时间: {stats['max_execution_time']:.2f}秒\n")
            f.write(f"中位数执行时间: {stats['median_execution_time']:.2f}秒\n\n")

            f.write("【风险评估矩阵】\n")
            f.write("-" * 80 + "\n")
            f.write(f"{'风险等级':<10} {'测试数量':<10} {'通过':<8} {'失败':<8} {'通过率':<10} {'占比':<10}\n")
            f.write("-" * 80 + "\n")
            for level in ["critical", "high", "medium", "low"]:
                count = stats['risk_distribution'][level]
                passed = stats['risk_status'][level]['passed']
                failed = stats['risk_status'][level]['failed']
                pass_rate = (passed / count * 100) if count > 0 else 0
                percentage = (count / stats['total_tests'] * 100) if stats['total_tests'] > 0 else 0
                f.write(f"{level:<10} {count:<10} {passed:<8} {failed:<8} {pass_rate:<10.1f} {percentage:<10.1f}\n")
            f.write("\n")

            # ATT&CK覆盖热力图
            f.write("【ATT&CK 覆盖热力图】\n")
            f.write("-" * 100 + "\n")
            f.write("热力图展示各战术下技术的测试覆盖情况。图例：[优秀>=90%] [良好70-89%] [一般50-69%] [较差25-49%] [严重<25%]\n")
            f.write("-" * 100 + "\n")

            heatmap_data = self.get_attack_heatmap_data()
            for tactic in heatmap_data['tactics']:
                tactic_info = heatmap_data['tactic_summary'].get(tactic, {})
                tech_list = sorted(
                    heatmap_data['technique_by_tactic'].get(tactic, []),
                    key=lambda x: x['pass_rate']
                )

                f.write(f"\n{tactic}:\n")
                f.write(f"  技术数: {tactic_info.get('technique_count', 0)} | ")
                f.write(f"测试数: {tactic_info.get('total_tests', 0)} | ")
                f.write(f"通过: {tactic_info.get('passed', 0)} | ")
                f.write(f"失败: {tactic_info.get('failed', 0)} | ")
                pass_rate = tactic_info.get('pass_rate', 0)
                if pass_rate >= 90:
                    status = "[优秀]"
                elif pass_rate >= 70:
                    status = "[良好]"
                elif pass_rate >= 50:
                    status = "[一般]"
                elif pass_rate >= 25:
                    status = "[较差]"
                else:
                    status = "[严重]"
                f.write(f"通过率: {pass_rate:.1f}% {status}\n")

                if tech_list:
                    f.write("  技术详情:\n")
                    for tech_data in tech_list[:8]:  # 最多显示8个技术
                        tech_pass = tech_data['pass_rate']
                        if tech_pass >= 90:
                            tech_status = "✓"
                        elif tech_pass >= 70:
                            tech_status = "○"
                        elif tech_pass >= 50:
                            tech_status = "△"
                        else:
                            tech_status = "✗"
                        f.write(f"    {tech_data['technique']}: {tech_pass:.0f}% {tech_status} ({tech_data['total']}次测试)\n")
            f.write("\n")

            f.write("【战术分类统计】\n")
            f.write("-" * 80 + "\n")
            f.write(f"{'战术类别':<15} {'测试次数':<10} {'技术数':<10} {'通过':<8} {'失败':<8} {'通过率':<10}\n")
            f.write("-" * 80 + "\n")
            for tactic, data in sorted(stats['tactic_stats'].items(), key=lambda x: x[1]['total'], reverse=True):
                f.write(f"{tactic:<15} {data['total']:<10} {data['technique_count']:<10} {data['passed']:<8} {data['failed']:<8} {data['pass_rate']:<10.1f}\n")
            f.write("\n")

            if stats['error_types']:
                f.write("【错误类型分析】\n")
                f.write("-" * 80 + "\n")
                total_errors = sum(stats['error_types'].values())
                for error_type, count in sorted(stats['error_types'].items(), key=lambda x: x[1], reverse=True):
                    percentage = (count / total_errors * 100) if total_errors > 0 else 0
                    f.write(f"{error_type}: {count}次 ({percentage:.1f}%)\n")
                f.write("\n")

            f.write("【详细技术分析】\n")
            f.write("-" * 100 + "\n")
            f.write(f"{'技术ID':<10} {'战术':<10} {'测试':<6} {'通过%':<8} {'平均(秒)':<10} {'最短':<8} {'最长':<8} {'稳定性':<8}\n")
            f.write("-" * 100 + "\n")
            for tech, analysis in sorted(detailed_analysis.items(), key=lambda x: x[1]['total_tests'], reverse=True):
                f.write(f"{tech:<10} {analysis['tactic']:<10} {analysis['total_tests']:<6} "
                       f"{analysis['pass_rate']:<8.1f} {analysis['avg_execution_time']:<10.2f} "
                       f"{analysis['min_execution_time']:<8.2f} {analysis['max_execution_time']:<8.2f} "
                       f"{analysis['stability']:<8}\n")
            f.write("\n")

            f.write("【详细测试结果】\n")
            f.write("-" * 80 + "\n")
            for result in self.results:
                f.write(f"\n技术ID: {result['technique']} ({result.get('tactic', '未知')})\n")
                f.write(f"描述: {result['description']}\n")
                f.write(f"文件: {result['filename']}\n")
                f.write(f"状态: {result['status']}\n")
                f.write(f"时间: {result['timestamp']}\n")
                f.write(f"执行时间: {result['execution_time']}秒\n")
                f.write(f"错误类型: {result.get('error_type', 'N/A')}\n")

                if result['commands']:
                    f.write("执行的命令:\n")
                    for cmd in result['commands'][:3]:
                        f.write(f"  - {cmd}\n")

                if result['indicators']:
                    f.write("检测指标:\n")
                    for indicator in result['indicators'][:3]:
                        f.write(f"  - {indicator}\n")

                if result['errors']:
                    f.write("错误信息:\n")
                    for error in result['errors'][:2]:
                        f.write(f"  - {error}\n")

                f.write("-" * 40 + "\n")

            f.write("\n【修复建议与行动计划】\n")
            f.write("=" * 80 + "\n")

            f.write("\n>> 立即行动（0-24小时）\n")
            f.write("-" * 80 + "\n")
            for item in remediation_plan['immediate']:
                f.write(f"[{item['priority']}] {item['title']}\n")
                f.write(f"  描述: {item['description']}\n\n")

            f.write("\n>> 短期改进（1-2周）\n")
            f.write("-" * 80 + "\n")
            for item in remediation_plan['short_term']:
                f.write(f"[{item['priority']}] {item['title']}\n")
                f.write(f"  描述: {item['description']}\n\n")

            f.write("\n>> 中期改进（1-3个月）\n")
            f.write("-" * 80 + "\n")
            for item in remediation_plan['medium_term']:
                f.write(f"[{item['priority']}] {item['title']}\n")
                f.write(f"  描述: {item['description']}\n\n")

            f.write("\n>> 长期优化（3个月以上）\n")
            f.write("-" * 80 + "\n")
            for item in remediation_plan['long_term']:
                f.write(f"[{item['priority']}] {item['title']}\n")
                f.write(f"  描述: {item['description']}\n\n")

        print(f"文本报告已生成: {output_file}")
        return output_file

    def generate_docx_report(self, output_file: str = None):
        """生成增强版Word格式的详细报告，具备详细修复指导和可读性"""
        if not DOCX_AVAILABLE:
            print("警告: python-docx 库未安装，无法生成Word报告")
            print("请运行: pip install python-docx")
            return None

        if not output_file:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_file = f"{self.report_dir}/security_validation_report_windows_{timestamp}.docx"

        stats = self.calculate_statistics()
        detailed_analysis = self.get_detailed_technique_analysis()
        remediation_plan = self.get_remediation_plan(stats)

        # 创建Word文档
        doc = Document()

        # 设置中文字体
        def set_chinese_font(run, font_name='SimSun', font_size=10.5, bold=False):
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
            run.font.bold = bold

        # XML字符清理函数
        def sanitize_xml(text):
            """清理XML不兼容字符"""
            if text is None:
                return ""
            if not isinstance(text, str):
                text = str(text)
            cleaned = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
            cleaned = cleaned.replace('\x00', '')
            return cleaned

        # 辅助函数：安全设置单元格文本
        def set_cell_text(cell, text):
            cell.text = sanitize_xml(text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run)

        # 辅助函数：添加带样式的段落
        def add_styled_para(doc, text, bold=False, color=None, font_size=10.5):
            para = doc.add_paragraph()
            run = para.add_run(sanitize_xml(text))
            run.bold = bold
            set_chinese_font(run, font_size=font_size)
            if color:
                run.font.color.rgb = color
            return para

        # 计算安全评分
        security_score = (stats['passed'] / stats['total_tests'] * 100) if stats['total_tests'] > 0 else 0
        score_grade = 'A' if security_score >= 90 else ('B' if security_score >= 80 else ('C' if security_score >= 70 else ('D' if security_score >= 60 else 'F')))
        score_color = RGBColor(0, 128, 0) if security_score >= 80 else (RGBColor(255, 165, 0) if security_score >= 60 else RGBColor(255, 0, 0))

        # 添加标题
        title = doc.add_heading('MITRE ATT&CK 主机安全验证测试报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            set_chinese_font(run, font_name='SimHei', font_size=22, bold=True)

        # 副标题
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('基于 Atomic Red Team 的攻防验证评估')
        set_chinese_font(run, font_name='SimHei', font_size=12)
        subtitle.add_run('\n')
        run = subtitle.add_run(f'报告生成时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        set_chinese_font(run, font_size=10)

        doc.add_paragraph()

        # 报告概述框（前言，不纳入目录）
        overview_title = doc.add_paragraph()
        overview_run = overview_title.add_run('报告概述')
        overview_run.bold = True
        overview_run.font.size = Pt(16)
        set_chinese_font(overview_run, font_name='SimHei', font_size=16, bold=True)

        overview_para = doc.add_paragraph()
        overview_text = (
            f"本报告基于 MITRE ATT&CK 框架对系统进行了全面的安全验证测试，共执行 {stats['total_tests']} 次测试，"
            f"涵盖 {stats['techniques_tested']} 种攻击技术。测试范围覆盖Windows 系统环境，使用 Atomic Red Team "
            f"测试用例模拟真实攻击场景，评估现有安全控制的有效性和检测能力。"
        )
        run = overview_para.add_run(overview_text)
        set_chinese_font(run)

        doc.add_page_break()

        # 一、执行摘要
        heading = doc.add_heading('一、执行摘要与安全评分', 1)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=16, bold=True)

        # 安全评分卡片
        score_para = doc.add_paragraph()
        score_para.add_run('整体安全评分: ').bold = True
        score_run = score_para.add_run(f'{security_score:.1f}%')
        score_run.font.color.rgb = score_color
        score_run.bold = True
        set_chinese_font(score_run, font_size=14)
        grade_run = score_para.add_run(f'  (等级: {score_grade})')
        grade_run.bold = True
        set_chinese_font(grade_run, font_size=12)

        # 评分解读
        score_interp = doc.add_paragraph()
        if security_score >= 90:
            interp_text = '安全状况优秀：系统安全控制完善，检测能力较强，建议保持现有安全策略并持续监控。'
        elif security_score >= 80:
            interp_text = '安全状况良好：大部分安全控制有效，存在少量改进空间，建议关注失败测试项。'
        elif security_score >= 70:
            interp_text = '安全状况一般：存在明显的安全控制缺口，建议优先修复高风险失败项。'
        elif security_score >= 60:
            interp_text = '安全状况较差：安全控制存在较多漏洞，需要系统性整改。'
        else:
            interp_text = '安全状况严重：系统存在重大安全风险，需要立即采取修复措施。'
        run = score_interp.add_run(interp_text)
        set_chinese_font(run)

        doc.add_paragraph()

        # 关键指标表格
        heading = doc.add_heading('1.1 关键测试指标', 2)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        table = doc.add_table(rows=9, cols=3)
        table.style = 'Light Grid Accent 1'

        metrics_data = [
            ('指标项', '数值', '说明'),
            ('总测试数', str(stats['total_tests']), '执行的测试用例总数'),
            ('通过测试', f"{stats['passed']} ({stats['pass_rate']:.1f}%)", '安全控制成功阻止的攻击'),
            ('失败测试', str(stats['failed']), '攻击成功，安全控制未生效'),
            ('错误测试', str(stats['errors']), '测试执行过程出错'),
            ('警告测试', str(stats['warnings']), '测试完成但存在异常'),
            ('测试技术数', f"{stats['techniques_tested']}种", '覆盖的MITRE ATT&CK技术'),
            ('平均执行时间', f"{stats['avg_execution_time']:.2f}秒", '单次测试平均耗时'),
            ('关键失败数', str(stats['risk_status']['critical']['failed'] + stats['risk_status']['high']['failed']), '严重和高风险失败总数')
        ]

        for i, (metric, value, desc) in enumerate(metrics_data):
            row = table.rows[i]
            set_cell_text(row.cells[0], metric)
            set_cell_text(row.cells[1], value)
            set_cell_text(row.cells[2], desc)
            if i == 0:  # 表头加粗
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

        doc.add_paragraph()

        # 关键发现
        heading = doc.add_heading('1.2 关键发现', 2)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        findings = []

        # 发现1: 总体通过率
        if stats['pass_rate'] >= 80:
            findings.append(('积极发现', f'整体通过率达到 {stats["pass_rate"]:.1f}%，安全控制有效性良好'))
        else:
            findings.append(('需关注', f'整体通过率仅 {stats["pass_rate"]:.1f}%，存在较大改进空间'))

        # 发现2: 关键风险
        critical_failures = stats['risk_status']['critical']['failed']
        high_failures = stats['risk_status']['high']['failed']
        if critical_failures > 0:
            findings.append(('严重风险', f'发现 {critical_failures} 个严重风险测试失败，需立即处理'))
        if high_failures > 0:
            findings.append(('高风险', f'发现 {high_failures} 个高风险测试失败，需优先修复'))

        # 发现3: 低通过率战术
        low_pass_tactics = [(t, d) for t, d in stats['tactic_stats'].items()
                           if d['pass_rate'] < 50 and d['total'] >= 3]
        if low_pass_tactics:
            tactic_names = '、'.join([t[0] for t in low_pass_tactics[:3]])
            findings.append(('防御盲区', f'{tactic_names} 等战术类别防御薄弱，通过率低于50%'))

        # 发现4: 错误类型
        if stats['error_types']:
            top_error = max(stats['error_types'].items(), key=lambda x: x[1])
            findings.append(('执行问题', f'最常见的错误类型为 {top_error[0]}，共 {top_error[1]} 次，影响测试有效性'))

        for finding_type, finding_text in findings:
            para = doc.add_paragraph(style='List Bullet')
            type_run = para.add_run(f'[{finding_type}] ')
            type_run.bold = True
            if finding_type == '严重风险':
                type_run.font.color.rgb = RGBColor(255, 0, 0)
            elif finding_type == '积极发现':
                type_run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                type_run.font.color.rgb = RGBColor(255, 165, 0)
            set_chinese_font(type_run)

            text_run = para.add_run(finding_text)
            set_chinese_font(text_run)

        doc.add_page_break()

        # 二、风险评估详细分析
        heading = doc.add_heading('二、风险评估详细分析', 1)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=16, bold=True)

        risk_desc = doc.add_paragraph()
        risk_desc_text = (
            '本节按风险等级对测试结果进行分类分析。风险等级基于MITRE ATT&CK技术对企业系统的潜在影响'
            '进行划分：严重风险表示可能导致系统完全沦陷的技术，高风险表示可能获取系统控制权的技术，'
            '中低风险表示影响相对有限的技术。'
        )
        run = risk_desc.add_run(risk_desc_text)
        set_chinese_font(run)

        doc.add_paragraph()

        # 风险分布表格
        heading = doc.add_heading('2.1 风险等级分布统计', 2)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        table = doc.add_table(rows=5, cols=7)
        table.style = 'Light Grid Accent 1'

        headers = ['风险等级', '测试数量', '通过数', '失败数', '通过率', '占比', '状态评估']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                set_chinese_font(run, bold=True)

        risk_levels = ['critical', 'high', 'medium', 'low']
        risk_names = {'critical': '严重', 'high': '高', 'medium': '中', 'low': '低'}
        risk_colors = {
            'critical': ('严重', RGBColor(255, 0, 0)),
            'high': ('高风险', RGBColor(255, 69, 0)),
            'medium': ('中等', RGBColor(255, 165, 0)),
            'low': ('较低', RGBColor(0, 128, 0))
        }

        for idx, level in enumerate(risk_levels, 1):
            row = table.rows[idx]
            count = stats['risk_distribution'][level]
            passed = stats['risk_status'][level]['passed']
            failed = stats['risk_status'][level]['failed']
            pass_rate = (passed / count * 100) if count > 0 else 0
            percentage = (count / stats['total_tests'] * 100) if stats['total_tests'] > 0 else 0

            if failed == 0:
                status_text = '良好'
            elif level == 'critical':
                status_text = '紧急处理'
            elif level == 'high':
                status_text = '需关注'
            else:
                status_text = '可控'

            set_cell_text(row.cells[0], risk_names[level])
            set_cell_text(row.cells[1], str(count))
            set_cell_text(row.cells[2], str(passed))
            set_cell_text(row.cells[3], str(failed))
            set_cell_text(row.cells[4], f'{pass_rate:.1f}%')
            set_cell_text(row.cells[5], f'{percentage:.1f}%')
            set_cell_text(row.cells[6], status_text)

        doc.add_paragraph()

        # 风险解读
        heading = doc.add_heading('2.2 风险评估解读', 2)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        if critical_failures > 0:
            add_styled_para(doc, f'⚠ 严重风险警报：发现 {critical_failures} 个严重风险测试失败。这些技术通常用于实现系统完全控制、数据窃取或服务中断。失败的测试表明系统对这些高级威胁缺乏有效防护，攻击者可能利用这些技术获取系统最高权限。', color=RGBColor(255, 0, 0))

        if high_failures > 0:
            add_styled_para(doc, f'⚠ 高风险警报：发现 {high_failures} 个高风险测试失败。这些技术常用于权限提升、横向移动和持久化。未能检测或阻止这些攻击可能导致攻击者在网络中扩散和长期潜伏。', color=RGBColor(255, 69, 0))

        # 添加风险缓解建议
        risk_mitigation = doc.add_paragraph()
        risk_mitigation.add_run('风险缓解建议：').bold = True
        for run in risk_mitigation.runs:
            set_chinese_font(run, bold=True)

        # 根据实际风险情况生成针对性建议
        risk_mitigation_detailed = []

        if critical_failures > 0:
            risk_mitigation_detailed.append({
                'title': '严重风险紧急处置',
                'priority': 'P0',
                'timeline': '立即（0-4小时）',
                'actions': [
                    '隔离受影响系统，限制潜在攻击面',
                    '检查是否存在实际入侵痕迹（日志审计、进程分析）',
                    '临时启用增强监控模式，收集更详细的审计数据',
                    '通知安全负责人和应急响应团队',
                    '准备回滚方案，确保业务连续性'
                ],
                'success_criteria': '严重风险测试通过率达到100%，无新增安全事件'
            })

        if high_failures > 0:
            risk_mitigation_detailed.append({
                'title': '高风险漏洞修复',
                'priority': 'P1',
                'timeline': '短期（0-24小时）',
                'actions': [
                    '分析失败测试的具体技术，了解攻击原理',
                    '评估现有安全控制为何未能检测/阻止该攻击',
                    '部署针对性的检测规则（如Sigma、YARA规则）',
                    '配置相应的阻断策略（如EDR阻断规则、防火墙策略）',
                    '验证修复效果，确保测试通过'
                ],
                'success_criteria': '高风险测试通过率达到90%以上'
            })

        risk_mitigation_detailed.extend([
            {
                'title': '纵深防御体系建设',
                'priority': 'P2',
                'timeline': '中期（1-2周）',
                'actions': [
                    '梳理现有安全控制体系，识别单点故障风险',
                    '在每个攻击阶段部署多层防护措施',
                    '建立预防-检测-响应的完整安全链条',
                    '定期测试各层防护的有效性',
                    '制定安全控制失效的应急响应预案'
                ],
                'success_criteria': '各战术类别均有至少两层防护措施'
            },
            {
                'title': '持续监控与威胁狩猎',
                'priority': 'P2',
                'timeline': '持续进行',
                'actions': [
                    '建立基于MITRE ATT&CK的检测覆盖矩阵',
                    '配置SIEM告警规则，关联多源日志',
                    '开展定期威胁狩猎活动，主动发现潜在威胁',
                    '建立威胁情报订阅机制，及时更新检测规则',
                    '定期复盘安全事件，优化检测逻辑'
                ],
                'success_criteria': '平均检测时间(MTTD)低于1小时'
            }
        ])

        for item in risk_mitigation_detailed:
            # 标题
            para = doc.add_paragraph()
            priority_run = para.add_run(f"[{item['priority']}] ")
            priority_run.bold = True
            if item['priority'] == 'P0':
                priority_run.font.color.rgb = RGBColor(255, 0, 0)
            elif item['priority'] == 'P1':
                priority_run.font.color.rgb = RGBColor(255, 69, 0)
            else:
                priority_run.font.color.rgb = RGBColor(52, 152, 219)
            set_chinese_font(priority_run)

            title_run = para.add_run(item['title'])
            title_run.bold = True
            set_chinese_font(title_run)

            # 时间线
            timeline_para = doc.add_paragraph()
            timeline_para.add_run('响应时限：').bold = True
            run = timeline_para.add_run(item['timeline'])
            set_chinese_font(run)

            # 行动步骤
            actions_para = doc.add_paragraph()
            actions_para.add_run('具体措施：').bold = True
            for run in actions_para.runs:
                set_chinese_font(run, bold=True)

            for idx, action in enumerate(item['actions'], 1):
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(f"{idx}. {action}")
                set_chinese_font(run)

            # 成功标准
            success_para = doc.add_paragraph()
            success_para.add_run('成功标准：').bold = True
            run = success_para.add_run(item['success_criteria'])
            set_chinese_font(run)

            doc.add_paragraph()

        doc.add_page_break()

        # ATT&CK覆盖热力图
        heading = doc.add_heading('三、ATT&CK 覆盖热力图', 1)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=16, bold=True)

        heatmap_intro = doc.add_paragraph()
        heatmap_intro_text = (
            '以下热力图展示了各MITRE ATT&CK战术下技术的测试覆盖情况和通过率。'
            '颜色编码：绿色表示高通过率（安全控制有效），黄色表示中等通过率，红色表示低通过率（存在安全风险）。'
        )
        run = heatmap_intro.add_run(heatmap_intro_text)
        set_chinese_font(run)

        doc.add_paragraph()

        # 获取热力图数据
        heatmap_data = self.get_attack_heatmap_data()

        # 创建热力图表格
        heading = doc.add_heading('3.1 战术-技术覆盖矩阵', 2)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        # 计算表格大小
        max_techs_per_tactic = max(
            len(heatmap_data['technique_by_tactic'].get(tactic, []))
            for tactic in heatmap_data['tactics']
        ) if heatmap_data['tactics'] else 0

        # 限制技术列数量
        tech_columns = min(max_techs_per_tactic, 8)
        total_cols = 6 + tech_columns  # 战术名 + 5个统计列 + 技术列

        row_count = len(heatmap_data['tactics']) + 1
        table = doc.add_table(rows=row_count, cols=total_cols)
        table.style = 'Light Grid Accent 1'

        # 表头
        headers = ['战术类别', '技术数', '测试数', '通过数', '失败数', '通过率']
        for i in range(tech_columns):
            headers.append(f'技术{i+1}')
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                set_chinese_font(run, bold=True)

        # 数据行
        for row_idx, tactic in enumerate(heatmap_data['tactics'], 1):
            tactic_info = heatmap_data['tactic_summary'].get(tactic, {})
            tech_list = sorted(
                heatmap_data['technique_by_tactic'].get(tactic, []),
                key=lambda x: x['pass_rate']
            )[:tech_columns]

            row = table.rows[row_idx]
            set_cell_text(row.cells[0], tactic)
            set_cell_text(row.cells[1], str(tactic_info.get('technique_count', 0)))
            set_cell_text(row.cells[2], str(tactic_info.get('total_tests', 0)))
            set_cell_text(row.cells[3], str(tactic_info.get('passed', 0)))
            set_cell_text(row.cells[4], str(tactic_info.get('failed', 0)))
            set_cell_text(row.cells[5], f"{tactic_info.get('pass_rate', 0):.1f}%")

            # 技术列
            for tech_idx, tech_data in enumerate(tech_list):
                cell = table.rows[row_idx].cells[6 + tech_idx]
                pass_rate = tech_data['pass_rate']
                tech_text = f"{tech_data['technique']}\n{pass_rate:.0f}%"
                set_cell_text(cell, tech_text)

            # 填充空白技术列
            for tech_idx in range(len(tech_list), tech_columns):
                set_cell_text(row.cells[6 + tech_idx], '-')

        doc.add_paragraph()

        # 热力图说明
        legend_para = doc.add_paragraph()
        legend_para.add_run('热力图颜色说明：').bold = True
        for run in legend_para.runs:
            set_chinese_font(run, bold=True)

        legend_items = [
            ('优秀 (≥90%)', '安全控制非常有效，测试通过率高'),
            ('良好 (70-89%)', '安全控制基本有效，存在少量改进空间'),
            ('一般 (50-69%)', '安全控制部分有效，需要加强'),
            ('较差 (25-49%)', '安全控制不足，需要重点关注'),
            ('严重 (<25%)', '安全控制失效，需要立即修复')
        ]

        for level, desc in legend_items:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(f'{level}: ').bold = True
            para.add_run(desc)
            for run in para.runs:
                set_chinese_font(run)

        doc.add_page_break()

        # 三、战术分类统计与分析
        heading = doc.add_heading('四、战术分类统计与分析', 1)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=16, bold=True)

        tactic_desc = doc.add_paragraph()
        tactic_desc_text = (
            '本节按MITRE ATT&CK战术类别（Tactic）对测试结果进行统计。每个战术代表攻击者达成特定目标'
            '的攻击阶段，如初始访问、执行、持久化、权限提升等。分析各战术类别的通过率有助于识别'
            '防御体系的薄弱环节。'
        )
        run = tactic_desc.add_run(tactic_desc_text)
        set_chinese_font(run)

        doc.add_paragraph()

        # 战术统计表格
        heading = doc.add_heading('4.1 战术类别详细统计', 2)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        tactic_count = len(stats['tactic_stats'])
        table = doc.add_table(rows=tactic_count + 1, cols=7)
        table.style = 'Light Grid Accent 1'

        headers = ['战术类别', '测试次数', '技术数', '通过数', '失败数', '通过率', '防御成熟度']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                set_chinese_font(run, bold=True)

        for idx, (tactic, data) in enumerate(sorted(stats['tactic_stats'].items(),
                                                     key=lambda x: x[1]['pass_rate']), 1):
            row = table.rows[idx]
            pass_rate = data['pass_rate']
            maturity = '成熟' if pass_rate >= 80 else ('基本' if pass_rate >= 60 else ('薄弱' if pass_rate >= 40 else '缺失'))

            set_cell_text(row.cells[0], tactic)
            set_cell_text(row.cells[1], str(data['total']))
            set_cell_text(row.cells[2], str(data['technique_count']))
            set_cell_text(row.cells[3], str(data['passed']))
            set_cell_text(row.cells[4], str(data['failed']))
            set_cell_text(row.cells[5], f"{pass_rate:.1f}%")
            set_cell_text(row.cells[6], maturity)

        doc.add_paragraph()

        # 战术分析解读
        if low_pass_tactics:
            heading = doc.add_heading('4.2 防御薄弱战术分析', 2)
            for run in heading.runs:
                set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

            for idx, (tactic, data) in enumerate(low_pass_tactics[:3], 1):
                tactic_heading = doc.add_heading(f'4.2.{idx} {tactic}', 3)
                for run in tactic_heading.runs:
                    set_chinese_font(run, font_name='SimHei', font_size=12, bold=True)

                para = doc.add_paragraph()
                analysis_text = (
                    f'该战术类别测试通过率为 {data["pass_rate"]:.1f}%（{data["passed"]}/{data["total"]}），'
                    f'表明系统对 {tactic} 阶段的攻击行为检测和防御能力较弱。'
                )
                run = para.add_run(analysis_text)
                set_chinese_font(run)

                # 具体建议
                suggestion_para = doc.add_paragraph()
                suggestion_para.add_run('改进建议：').bold = True
                for run in suggestion_para.runs:
                    set_chinese_font(run, bold=True)

                # 详细的战术改进建议
                tactic_recommendations = {
                    '初始访问': {
                        'description': '初始访问是攻击者进入企业网络的第一步，主要通过钓鱼邮件、漏洞利用或供应链攻击实现。',
                        'recommendations': [
                            {'action': '部署邮件安全网关', 'detail': '配置SPF/DKIM/DMARC，启用高级威胁防护(ATP)，过滤恶意附件和链接', 'tool': 'Microsoft Defender for Office 365 / Proofpoint'},
                            {'action': '实施应用白名单', 'detail': '限制只能运行经批准的应用程序，阻止未授权代码执行', 'tool': 'AppLocker / Carbon Black'},
                            {'action': '加强用户安全意识', 'detail': '定期开展钓鱼演练，培训员工识别可疑邮件和链接', 'tool': 'KnowBe4 / PhishMe'},
                            {'action': '加固面向公众的应用', 'detail': '定期漏洞扫描和渗透测试，及时修补已知漏洞', 'tool': 'Qualys / Nessus'},
                            {'action': '实施网络分段', 'detail': '隔离DMZ区域，限制外部访问内部资源的路径', 'tool': '防火墙策略 / 微分段'}
                        ]
                    },
                    '执行': {
                        'description': '执行阶段攻击者尝试在目标系统上运行恶意代码，常用手段包括命令行、脚本、计划任务等。',
                        'recommendations': [
                            {'action': '部署EDR/EPP解决方案', 'detail': '实时监控进程行为，检测和阻止可疑执行', 'tool': 'CrowdStrike / SentinelOne / Defender for Endpoint'},
                            {'action': '约束PowerShell使用', 'detail': '启用约束语言模式，记录脚本块日志，限制远程执行', 'tool': 'PowerShell策略 / AMSI'},
                            {'action': '监控脚本执行', 'detail': '记录所有脚本执行行为，检测混淆和可疑模式', 'tool': 'Sysmon / Windows Script Logging'},
                            {'action': '禁用不必要的执行载体', 'detail': '限制wscript、cscript、mshta等工具的使用', 'tool': 'AppLocker / 组策略'},
                            {'action': '实施运行时保护', 'detail': '启用DEP/ASLR，部署反漏洞利用保护', 'tool': 'EMET / Windows Exploit Guard'}
                        ]
                    },
                    '持久化': {
                        'description': '持久化使攻击者能够在系统中保持长期存在，即使系统重启也不会丢失访问权限。',
                        'recommendations': [
                            {'action': '监控启动项变更', 'detail': '监控注册表Run键、启动文件夹、服务的变更', 'tool': 'Sysmon / EDR规则'},
                            {'action': '审计计划任务', 'detail': '监控计划任务的创建和修改，检测异常任务', 'tool': 'Windows事件日志 / SIEM规则'},
                            {'action': '审查服务配置', 'detail': '定期审计服务列表，检测异常服务安装', 'tool': 'SCCM / 服务审计脚本'},
                            {'action': '监控WMI事件订阅', 'detail': '检测异常WMI事件消费者和过滤器', 'tool': 'WMI审计工具 / Sysmon'},
                            {'action': '实施配置基线', 'detail': '建立系统配置基线，定期检查配置漂移', 'tool': 'SCCM / DSC'}
                        ]
                    },
                    '权限提升': {
                        'description': '权限提升使攻击者能够获取更高权限（如SYSTEM/ROOT），从而执行更敏感的操作。',
                        'recommendations': [
                            {'action': '实施最小权限原则', 'detail': '移除不必要的本地管理员权限，使用标准用户账户', 'tool': 'LAPS / PAM解决方案'},
                            {'action': '及时安装安全补丁', 'detail': '优先修补提权相关漏洞，建立补丁管理流程', 'tool': 'WSUS / SCCM / 自动化补丁管理'},
                            {'action': '监控特权操作', 'detail': '记录所有特权账户活动，检测异常提权行为', 'tool': 'PAM / SIEM规则'},
                            {'action': '启用UAC并配置合理级别', 'detail': '防止静默提权，要求管理员确认敏感操作', 'tool': '组策略 / UAC配置'},
                            {'action': '审计令牌和票据', 'detail': '监控令牌模拟和Kerberos票据活动', 'tool': 'RDPASSWRD检测 / Kerberos审计'}
                        ]
                    },
                    '防御规避': {
                        'description': '防御规避技术帮助攻击者隐藏其活动，绕过安全检测和日志记录。',
                        'recommendations': [
                            {'action': '增强日志记录', 'detail': '启用详细审计策略，记录命令行参数和脚本内容', 'tool': 'Windows高级审计 / Sysmon'},
                            {'action': '部署行为分析', 'detail': '使用ML/AI检测异常行为模式，识别未知威胁', 'tool': 'UEBA / EDR行为分析'},
                            {'action': '监控日志清除行为', 'detail': '告警任何日志清除或修改尝试', 'tool': 'SIEM规则 / Sysmon事件1100-1102'},
                            {'action': '实施日志集中存储', 'detail': '将日志转发到中央服务器，防止本地篡改', 'tool': 'SIEM / 日志收集器'},
                            {'action': '启用反调试保护', 'detail': '防止攻击者调试和逆向安全软件', 'tool': '安全软件自保护功能'}
                        ]
                    },
                    '凭据访问': {
                        'description': '凭据访问使攻击者能够获取账户凭据，用于横向移动或权限提升。',
                        'recommendations': [
                            {'action': '实施Credential Guard', 'detail': '使用虚拟化保护凭据，防止LSASS内存转储', 'tool': 'Windows Defender Credential Guard'},
                            {'action': '限制LSASS访问', 'detail': '配置LSASS保护模式，阻止非授权进程访问', 'tool': 'PPL / LSASS保护策略'},
                            {'action': '监控凭据转储行为', 'detail': '检测Mimikatz等工具的特征和行为', 'tool': 'EDR规则 / Sigma规则'},
                            {'action': '启用多因素认证', 'detail': '为所有敏感账户启用MFA，减少凭据泄露影响', 'tool': 'Azure AD MFA / Duo'},
                            {'action': '实施密码策略', 'detail': '强制复杂密码，定期轮换，防止密码复用', 'tool': '密码策略 / PAM'}
                        ]
                    },
                    '发现': {
                        'description': '发现阶段攻击者收集系统和网络信息，为后续攻击做准备。',
                        'recommendations': [
                            {'action': '监控侦察命令', 'detail': '记录和告警系统信息收集命令（whoami、net、systeminfo等）', 'tool': 'Sysmon / SIEM规则'},
                            {'action': '限制信息暴露', 'detail': '配置系统最小化信息返回，减少可收集的数据', 'tool': '组策略 / 系统配置'},
                            {'action': '检测异常查询频率', 'detail': '识别短时间内大量侦察行为的模式', 'tool': 'UEBA / 行为分析'},
                            {'action': '实施网络分段', 'detail': '限制系统间的信息可见性，增加侦察难度', 'tool': '网络分段 / 微隔离'},
                            {'action': '蜜罐和诱饵部署', 'detail': '部署虚假资源吸引侦察行为', 'tool': '蜜罐解决方案'}
                        ]
                    },
                    '横向移动': {
                        'description': '横向移动使攻击者能够在网络内部扩散，访问更多系统和数据。',
                        'recommendations': [
                            {'action': '实施网络微分段', 'detail': '基于应用和工作负载隔离网络，限制横向通信', 'tool': 'NSX / 微隔离平台'},
                            {'action': '限制远程访问协议', 'detail': '严格管控RDP、WinRM、SMB等协议的使用', 'tool': '防火墙策略 / 组策略'},
                            {'action': '监控远程访问', 'detail': '记录所有远程连接，检测异常登录模式', 'tool': 'SIEM / EDR'},
                            {'action': '实施跳跃服务器', 'detail': '通过堡垒机管理远程访问，集中审计', 'tool': '堡垒机 / PAM'},
                            {'action': '检测票据传递攻击', 'detail': '监控Kerberos票据异常，检测PtT/PtH攻击', 'tool': 'Kerberos审计 / EDR规则'}
                        ]
                    },
                    '数据收集': {
                        'description': '数据收集阶段攻击者搜索和收集感兴趣的数据，为外泄做准备。',
                        'recommendations': [
                            {'action': '数据分类和标记', 'detail': '对敏感数据进行分类，实施差异化保护', 'tool': 'DLP解决方案 / 信息保护'},
                            {'action': '监控数据访问', 'detail': '记录敏感文件的访问、复制和修改行为', 'tool': 'DLP / 文件审计'},
                            {'action': '检测数据压缩行为', 'detail': '监控大量压缩操作，可能是数据收集的迹象', 'tool': 'EDR / 行为分析'},
                            {'action': '限制数据存储访问', 'detail': '实施最小权限访问，审计共享权限', 'tool': '文件服务器审计 / ACL管理'},
                            {'action': '部署数据防泄漏', 'detail': '监控和阻止敏感数据离开受控环境', 'tool': 'DLP解决方案'}
                        ]
                    },
                    '数据外泄': {
                        'description': '数据外泄是攻击者将收集的数据传输到外部受控服务器的过程。',
                        'recommendations': [
                            {'action': '监控出站流量', 'detail': '检测异常的数据传输行为和未知外联', 'tool': 'NGFW / Proxy日志分析'},
                            {'action': '限制云存储访问', 'detail': '管控对公共云存储服务的访问', 'tool': '云访问安全代理(CASB)'},
                            {'action': '检测DNS隧道', 'detail': '分析DNS流量，识别异常DNS请求模式', 'tool': 'DNS安全网关 / SIEM规则'},
                            {'action': '实施DLP策略', 'detail': '阻止敏感数据通过邮件、网盘等渠道外泄', 'tool': 'DLP解决方案'},
                            {'action': '建立流量基线', 'detail': '建立正常流量模式，识别异常数据传输', 'tool': 'NTA / NDR解决方案'}
                        ]
                    },
                    '命令与控制': {
                        'description': '命令与控制(C2)使攻击者能够与受感染系统通信，发送指令和接收数据。',
                        'recommendations': [
                            {'action': '检测异常外联', 'detail': '识别与未知IP/域名的通信，检测信标行为', 'tool': 'NTA / 威胁情报平台'},
                            {'action': '阻断已知恶意基础设施', 'detail': '使用威胁情报自动阻断已知C2地址', 'tool': '防火墙 / DNS过滤'},
                            {'action': '分析加密流量', 'detail': '检查SSL/TLS流量，识别恶意加密通信', 'tool': 'SSL检测代理 / EDR'},
                            {'action': '监控DNS请求', 'detail': '检测DGA域名和可疑DNS查询', 'tool': 'DNS安全解决方案'},
                            {'action': '实施网络隔离', 'detail': '限制系统与互联网的直接通信', 'tool': '网络分段 / Proxy'}
                        ]
                    }
                }

                # 获取当前战术的详细建议
                tactic_key_map = {
                    '初始访问': '初始访问', '执行': '执行', '持久化': '持久化',
                    '权限提升': '权限提升', '防御规避': '防御规避', '凭据访问': '凭据访问',
                    '发现': '发现', '横向移动': '横向移动', '数据收集': '数据收集',
                    '数据外泄': '数据外泄', '命令与控制': '命令与控制',
                    'Initial Access': '初始访问', 'Execution': '执行', 'Persistence': '持久化',
                    'Privilege Escalation': '权限提升', 'Defense Evasion': '防御规避',
                    'Credential Access': '凭据访问', 'Discovery': '发现', 'Lateral Movement': '横向移动',
                    'Collection': '数据收集', 'Exfiltration': '数据外泄', 'Command and Control': '命令与控制'
                }

                tactic_key = tactic_key_map.get(tactic, tactic)
                rec_data = tactic_recommendations.get(tactic_key, None)

                if rec_data:
                    # 描述
                    desc_para = doc.add_paragraph()
                    run = desc_para.add_run(rec_data['description'])
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(100, 100, 100)
                    set_chinese_font(run)

                    # 详细建议表格
                    table = doc.add_table(rows=len(rec_data['recommendations']) + 1, cols=3)
                    table.style = 'Light Grid Accent 1'

                    headers = ['措施', '详细说明', '推荐工具']
                    for i, header in enumerate(headers):
                        cell = table.rows[0].cells[i]
                        cell.text = header
                        for run in cell.paragraphs[0].runs:
                            set_chinese_font(run, bold=True)

                    for idx, rec in enumerate(rec_data['recommendations'], 1):
                        row = table.rows[idx]
                        set_cell_text(row.cells[0], rec['action'])
                        set_cell_text(row.cells[1], rec['detail'])
                        set_cell_text(row.cells[2], rec['tool'])
                else:
                    # 默认建议
                    default_suggestions = [
                        '部署针对该战术的检测规则',
                        '加强相关行为的监控和审计',
                        '参考MITRE ATT&CK官方缓解措施',
                        '定期测试安全控制的有效性'
                    ]
                    for suggestion in default_suggestions:
                        para = doc.add_paragraph(style='List Bullet')
                        run = para.add_run(suggestion)
                        set_chinese_font(run)

        doc.add_page_break()

        # 四、错误类型分析
        if stats['error_types']:
            heading = doc.add_heading('五、错误类型与执行问题分析', 1)
            for run in heading.runs:
                set_chinese_font(run, font_name='SimHei', font_size=16, bold=True)

            error_desc = doc.add_paragraph()
            error_desc_text = (
                '测试执行过程中遇到的错误会影响评估的完整性。本节分析错误类型和频率，'
                '帮助识别测试环境配置问题和测试用例兼容性问题。'
            )
            run = error_desc.add_run(error_desc_text)
            set_chinese_font(run)

            doc.add_paragraph()

            # 错误统计表格
            heading = doc.add_heading('5.1 错误类型统计', 2)
            for run in heading.runs:
                set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

            total_errors = sum(stats['error_types'].values())
            error_count = len(stats['error_types'])
            table = doc.add_table(rows=error_count + 1, cols=4)
            table.style = 'Light Grid Accent 1'

            headers = ['错误类型', '出现次数', '占比', '影响程度']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                for run in cell.paragraphs[0].runs:
                    set_chinese_font(run, bold=True)

            for idx, (error_type, count) in enumerate(sorted(stats['error_types'].items(),
                                                              key=lambda x: x[1], reverse=True), 1):
                row = table.rows[idx]
                percentage = (count / total_errors * 100) if total_errors > 0 else 0
                impact = '高' if percentage > 50 else ('中' if percentage > 20 else '低')

                set_cell_text(row.cells[0], error_type)
                set_cell_text(row.cells[1], str(count))
                set_cell_text(row.cells[2], f'{percentage:.1f}%')
                set_cell_text(row.cells[3], impact)

            doc.add_paragraph()

            # 错误解决指南
            heading = doc.add_heading('5.2 常见错误解决指南', 2)
            for run in heading.runs:
                set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

            error_solutions = {
                'timeout': {
                    'cause': '测试执行超时，可能由于系统性能不足或测试用例执行时间过长',
                    'solution': ['调整测试超时设置（增加超时时间）', '优化被测系统性能', '检查网络连接稳定性', '在非业务高峰期执行测试']
                },
                'connection_refused': {
                    'cause': '无法连接到目标服务，可能服务未启动或网络不通',
                    'solution': ['验证相关服务运行状态', '检查防火墙规则和网络连通性', '确认端口配置正确', '检查SELinux/AppArmor策略']
                },
                'permission_denied': {
                    'cause': '执行权限不足，无法完成测试操作',
                    'solution': ['检查执行用户权限', '验证文件/目录权限设置', '确认SELinux/AppArmor策略允许', '使用sudo或以root身份执行']
                },
                'command_not_found': {
                    'cause': '测试依赖的命令或工具未安装',
                    'solution': ['安装缺失的依赖工具', '验证PATH环境变量配置', '确认工具版本兼容性', '建立标准化测试环境']
                }
            }

            for error_type, info in error_solutions.items():
                if error_type in [e.lower().replace(' ', '_') for e in stats['error_types'].keys()]:
                    para = doc.add_paragraph()
                    para.add_run(f'{error_type}: ').bold = True
                    for run in para.runs:
                        set_chinese_font(run, bold=True)

                    cause_para = doc.add_paragraph()
                    cause_para.add_run('原因: ').bold = True
                    run = cause_para.add_run(info['cause'])
                    set_chinese_font(run)

                    solution_para = doc.add_paragraph()
                    solution_para.add_run('解决方案: ').bold = True
                    for run in solution_para.runs:
                        set_chinese_font(run, bold=True)

                    for solution in info['solution']:
                        para = doc.add_paragraph(style='List Bullet')
                        run = para.add_run(solution)
                        set_chinese_font(run)

                    doc.add_paragraph()

            doc.add_page_break()

        # 五、修复建议与行动计划
        heading = doc.add_heading('六、修复建议与行动计划', 1)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=16, bold=True)

        # 修复建议说明
        remediation_intro = doc.add_paragraph()
        remediation_intro_text = (
            '本节基于测试结果提供分级修复建议，每项建议包含具体实施步骤、责任人、验证标准、'
            '预计工作量和业务影响评估。建议按照优先级顺序执行，确保关键风险得到及时处理。'
        )
        run = remediation_intro.add_run(remediation_intro_text)
        set_chinese_font(run)

        doc.add_paragraph()

        # 优先级说明
        priority_heading = doc.add_heading('6.1 优先级说明', 2)
        for run in priority_heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        priority_table = doc.add_table(rows=5, cols=3)
        priority_table.style = 'Light Grid Accent 1'

        priority_data = [
            ('优先级', '响应时间', '说明'),
            ('P0 - 紧急', '0-4小时', '严重安全风险，可能导致系统完全沦陷，需立即处理'),
            ('P1 - 高', '0-24小时', '高风险安全问题，对系统安全有重大影响，需优先处理'),
            ('P2 - 中', '1-2周', '中等风险问题，有计划地修复'),
            ('P3 - 低', '1-3个月', '长期改进项目，系统性提升安全能力')
        ]

        for i, (priority, time, desc) in enumerate(priority_data):
            row = priority_table.rows[i]
            set_cell_text(row.cells[0], priority)
            set_cell_text(row.cells[1], time)
            set_cell_text(row.cells[2], desc)
            if i == 0:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

        doc.add_paragraph()

        # 立即行动
        heading = doc.add_heading('6.2 立即行动（0-24小时）', 2)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        if remediation_plan['immediate']:
            for item in remediation_plan['immediate']:
                # 标题
                para = doc.add_paragraph()
                para.add_run(f"[{item['priority']}] {item['title']}").bold = True
                for run in para.runs:
                    set_chinese_font(run, bold=True)

                # 描述
                desc_para = doc.add_paragraph()
                desc_para.add_run('问题描述: ').bold = True
                run = desc_para.add_run(sanitize_xml(item['description']))
                set_chinese_font(run)

                # 业务影响
                if 'impact' in item:
                    impact_para = doc.add_paragraph()
                    impact_para.add_run('业务影响: ').bold = True
                    run = impact_para.add_run(sanitize_xml(item['impact']))
                    set_chinese_font(run)
                    impact_para.runs[0].font.color.rgb = RGBColor(255, 69, 0)

                # 实施步骤
                if 'action_steps' in item:
                    steps_para = doc.add_paragraph()
                    steps_para.add_run('实施步骤:').bold = True
                    for run in steps_para.runs:
                        set_chinese_font(run, bold=True)

                    for idx, step in enumerate(item['action_steps'], 1):
                        para = doc.add_paragraph(style='List Bullet')
                        run = para.add_run(f"{idx}. {sanitize_xml(step)}")
                        set_chinese_font(run)

                # 负责人
                if 'owners' in item:
                    owner_para = doc.add_paragraph()
                    owner_para.add_run('责任团队: ').bold = True
                    run = owner_para.add_run(', '.join(item['owners']))
                    set_chinese_font(run)

                # 验证标准
                if 'verification' in item:
                    verify_para = doc.add_paragraph()
                    verify_para.add_run('验证标准: ').bold = True
                    run = verify_para.add_run(sanitize_xml(item['verification']))
                    set_chinese_font(run)

                # 预计工作量
                if 'estimated_effort' in item:
                    effort_para = doc.add_paragraph()
                    effort_para.add_run('预计工作量: ').bold = True
                    run = effort_para.add_run(item['estimated_effort'])
                    set_chinese_font(run)

                # 涉及的技术
                if 'techniques' in item and item['techniques']:
                    tech_para = doc.add_paragraph()
                    tech_para.add_run('涉及技术ID: ').bold = True
                    run = tech_para.add_run(', '.join(item['techniques'][:10]))
                    set_chinese_font(run)

                doc.add_paragraph()
        else:
            para = doc.add_paragraph()
            run = para.add_run('✓ 未发现需要立即处理的高风险问题')
            run.font.color.rgb = RGBColor(0, 128, 0)
            set_chinese_font(run)
            doc.add_paragraph()

        doc.add_page_break()

        # 短期改进
        heading = doc.add_heading('6.3 短期改进（1-2周）', 2)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        for item in remediation_plan['short_term']:
            para = doc.add_paragraph()
            para.add_run(f"[{item['priority']}] {item['title']}").bold = True
            for run in para.runs:
                set_chinese_font(run, bold=True)

            desc_para = doc.add_paragraph()
            desc_para.add_run('目标: ').bold = True
            run = desc_para.add_run(sanitize_xml(item['description']))
            set_chinese_font(run)

            if 'impact' in item:
                impact_para = doc.add_paragraph()
                impact_para.add_run('预期收益: ').bold = True
                run = impact_para.add_run(sanitize_xml(item['impact']))
                set_chinese_font(run)

            if 'action_steps' in item:
                steps_para = doc.add_paragraph()
                steps_para.add_run('关键行动:').bold = True
                for run in steps_para.runs:
                    set_chinese_font(run, bold=True)

                for idx, step in enumerate(item['action_steps'][:5], 1):
                    para = doc.add_paragraph(style='List Bullet')
                    run = para.add_run(f"{idx}. {sanitize_xml(step)}")
                    set_chinese_font(run)

            if 'owners' in item:
                owner_para = doc.add_paragraph()
                owner_para.add_run('负责团队: ').bold = True
                run = owner_para.add_run(', '.join(item['owners']))
                set_chinese_font(run)

            if 'verification' in item:
                verify_para = doc.add_paragraph()
                verify_para.add_run('完成标准: ').bold = True
                run = verify_para.add_run(sanitize_xml(item['verification']))
                set_chinese_font(run)

            if 'estimated_effort' in item:
                effort_para = doc.add_paragraph()
                effort_para.add_run('预计周期: ').bold = True
                run = effort_para.add_run(item['estimated_effort'])
                set_chinese_font(run)

            doc.add_paragraph()

        doc.add_page_break()

        # 中期改进
        heading = doc.add_heading('6.4 中期改进（1-3个月）', 2)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        for item in remediation_plan['medium_term']:
            para = doc.add_paragraph()
            para.add_run(f"[{item['priority']}] {item['title']}").bold = True
            for run in para.runs:
                set_chinese_font(run, bold=True)

            desc_para = doc.add_paragraph()
            desc_para.add_run('战略目标: ').bold = True
            run = desc_para.add_run(sanitize_xml(item['description']))
            set_chinese_font(run)

            if 'impact' in item:
                impact_para = doc.add_paragraph()
                impact_para.add_run('战略价值: ').bold = True
                run = impact_para.add_run(sanitize_xml(item['impact']))
                set_chinese_font(run)

            if 'action_steps' in item:
                steps_para = doc.add_paragraph()
                steps_para.add_run('实施路线:').bold = True
                for run in steps_para.runs:
                    set_chinese_font(run, bold=True)

                for idx, step in enumerate(item['action_steps'][:4], 1):
                    para = doc.add_paragraph(style='List Bullet')
                    run = para.add_run(f"{idx}. {sanitize_xml(step)}")
                    set_chinese_font(run)

            if 'owners' in item:
                owner_para = doc.add_paragraph()
                owner_para.add_run('负责部门: ').bold = True
                run = owner_para.add_run(', '.join(item['owners']))
                set_chinese_font(run)

            if 'verification' in item:
                verify_para = doc.add_paragraph()
                verify_para.add_run('成功指标: ').bold = True
                run = verify_para.add_run(sanitize_xml(item['verification']))
                set_chinese_font(run)

            if 'estimated_effort' in item:
                effort_para = doc.add_paragraph()
                effort_para.add_run('预计周期: ').bold = True
                run = effort_para.add_run(item['estimated_effort'])
                set_chinese_font(run)

            doc.add_paragraph()

        doc.add_page_break()

        # 长期优化
        heading = doc.add_heading('6.5 长期优化（3个月以上）', 2)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        for item in remediation_plan['long_term']:
            para = doc.add_paragraph()
            para.add_run(f"[{item['priority']}] {item['title']}").bold = True
            for run in para.runs:
                set_chinese_font(run, bold=True)

            desc_para = doc.add_paragraph()
            desc_para.add_run('愿景目标: ').bold = True
            run = desc_para.add_run(sanitize_xml(item['description']))
            set_chinese_font(run)

            if 'impact' in item:
                impact_para = doc.add_paragraph()
                impact_para.add_run('长远影响: ').bold = True
                run = impact_para.add_run(sanitize_xml(item['impact']))
                set_chinese_font(run)

            if 'action_steps' in item:
                steps_para = doc.add_paragraph()
                steps_para.add_run('建设路径:').bold = True
                for run in steps_para.runs:
                    set_chinese_font(run, bold=True)

                for idx, step in enumerate(item['action_steps'][:4], 1):
                    para = doc.add_paragraph(style='List Bullet')
                    run = para.add_run(f"{idx}. {sanitize_xml(step)}")
                    set_chinese_font(run)

            if 'owners' in item:
                owner_para = doc.add_paragraph()
                owner_para.add_run('主导部门: ').bold = True
                run = owner_para.add_run(', '.join(item['owners']))
                set_chinese_font(run)

            if 'verification' in item:
                verify_para = doc.add_paragraph()
                verify_para.add_run('成熟度指标: ').bold = True
                run = verify_para.add_run(sanitize_xml(item['verification']))
                set_chinese_font(run)

            doc.add_paragraph()

        doc.add_page_break()

        # 六、技术详细分析
        heading = doc.add_heading('七、技术详细分析', 1)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=16, bold=True)

        tech_intro = doc.add_paragraph()
        tech_intro_text = (
            '本节提供每种测试技术的详细分析，包括测试统计、稳定性评估和风险等级。'
            '数据按测试次数降序排列，帮助识别高频使用和高风险的技术。'
        )
        run = tech_intro.add_run(tech_intro_text)
        set_chinese_font(run)

        doc.add_paragraph()

        # 技术统计表格
        heading = doc.add_heading('7.1 技术测试统计表', 2)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        tech_count = len(detailed_analysis)
        table = doc.add_table(rows=min(tech_count + 1, 31), cols=8)
        table.style = 'Light Grid Accent 1'

        headers = ['技术ID', '战术类别', '测试次数', '通过率', '平均时间(秒)', '时间范围', '稳定性', '风险等级']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                set_chinese_font(run, bold=True)

        for idx, (tech, analysis) in enumerate(sorted(detailed_analysis.items(),
                                                       key=lambda x: x[1]['total_tests'], reverse=True)[:30], 1):
            row = table.rows[idx]
            time_range = f"{analysis['min_execution_time']:.2f}-{analysis['max_execution_time']:.2f}"

            set_cell_text(row.cells[0], tech)
            set_cell_text(row.cells[1], analysis['tactic'])
            set_cell_text(row.cells[2], str(analysis['total_tests']))
            set_cell_text(row.cells[3], f"{analysis['pass_rate']:.1f}%")
            set_cell_text(row.cells[4], f"{analysis['avg_execution_time']:.2f}")
            set_cell_text(row.cells[5], time_range)
            set_cell_text(row.cells[6], analysis['stability'])
            set_cell_text(row.cells[7], analysis['risk_level'])

        doc.add_paragraph()

        # 稳定性说明
        stability_para = doc.add_paragraph()
        stability_para.add_run('稳定性说明: ').bold = True
        for run in stability_para.runs:
            set_chinese_font(run, bold=True)

        stability_items = [
            '稳定 - 多次测试结果一致，安全控制可靠',
            '波动 - 测试结果不一致，存在间歇性问题',
            '不稳定 - 多次测试失败，安全控制未生效'
        ]

        for item in stability_items:
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(item)
            set_chinese_font(run)

        doc.add_page_break()

        # 七、详细测试结果
        heading = doc.add_heading('八、详细测试结果记录', 1)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=16, bold=True)

        result_intro = doc.add_paragraph()
        result_intro_text = (
            f'本节展示前50条详细测试结果记录。每条记录包含测试技术、执行状态、命令详情和错误信息。'
            '这些详细数据可用于深入分析具体失败原因和制定针对性修复措施。'
        )
        run = result_intro.add_run(result_intro_text)
        set_chinese_font(run)

        doc.add_paragraph()

        for result in self.results[:50]:
            para = doc.add_paragraph()
            para.add_run("技术ID: ").bold = True
            para.add_run(sanitize_xml(f"{result['technique']} ({result.get('tactic', '未知')})\n"))
            para.add_run("描述: ").bold = True
            para.add_run(sanitize_xml(f"{result['description']}\n"))
            para.add_run("状态: ").bold = True
            status_run = para.add_run(sanitize_xml(f"{result['status']}\n"))
            if result['status'] == 'passed':
                status_run.font.color.rgb = RGBColor(0, 128, 0)
            elif result['status'] in ['failed', 'error']:
                status_run.font.color.rgb = RGBColor(255, 0, 0)
            elif result['status'] == 'warning':
                status_run.font.color.rgb = RGBColor(255, 165, 0)
            para.add_run("时间: ").bold = True
            para.add_run(sanitize_xml(f"{result['timestamp']}\n"))
            para.add_run("执行时间: ").bold = True
            para.add_run(sanitize_xml(f"{result['execution_time']}秒" if result['execution_time'] else "N/A"))
            para.add_run("\n日志文件: ").bold = True
            para.add_run(sanitize_xml(result.get('filename', 'N/A')))

            for run in para.runs:
                set_chinese_font(run)

            if result['commands']:
                cmd_para = doc.add_paragraph()
                cmd_para.add_run("执行的命令:").bold = True
                for run in cmd_para.runs:
                    set_chinese_font(run)
                for cmd in result['commands'][:2]:
                    cmd_text = sanitize_xml(f"  - {cmd[:120]}...")
                    para = doc.add_paragraph(cmd_text, style='List Bullet')
                    for run in para.runs:
                        set_chinese_font(run)
                        run.font.size = Pt(9)

            if result['indicators']:
                ind_para = doc.add_paragraph()
                ind_para.add_run("检测指标:").bold = True
                for run in ind_para.runs:
                    set_chinese_font(run)
                for indicator in result['indicators'][:2]:
                    ind_text = sanitize_xml(f"  - {indicator[:120]}...")
                    para = doc.add_paragraph(ind_text, style='List Bullet')
                    for run in para.runs:
                        set_chinese_font(run)
                        run.font.size = Pt(9)

            if result['errors']:
                err_para = doc.add_paragraph()
                err_para.add_run("错误信息:").bold = True
                for run in err_para.runs:
                    set_chinese_font(run)
                for error in result['errors'][:2]:
                    err_text = sanitize_xml(f"  - {error[:120]}...")
                    para = doc.add_paragraph(err_text, style='List Bullet')
                    for run in para.runs:
                        set_chinese_font(run)
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.font.size = Pt(9)

            doc.add_paragraph("-" * 50)

        # 九、总结
        doc.add_page_break()
        heading = doc.add_heading('九、总结', 1)
        for run in heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=16, bold=True)

        # 总结性段落
        summary_para = doc.add_paragraph()
        summary_text = (
            f'本报告基于 {stats["total_tests"]} 次MITRE ATT&CK安全验证测试，全面评估了系统的安全防护能力。'
            f'测试覆盖了 {stats["techniques_tested"]} 种攻击技术，涉及 {len(stats["tactic_stats"])} 个战术类别。'
            f'整体安全评分为 {security_score:.1f}%（等级{score_grade}），'
        )
        if security_score >= 80:
            summary_text += '表明系统安全防护体系较为完善，具备较强的威胁检测和防御能力。'
        elif security_score >= 60:
            summary_text += '表明系统安全防护体系基本达标，但存在明显改进空间。'
        else:
            summary_text += '表明系统安全防护体系存在较大缺口，需要重点关注和整改。'

        run = summary_para.add_run(summary_text)
        set_chinese_font(run)

        # 行动指引
        action_guide = doc.add_paragraph()
        guide_text = f'''
详细的修复建议和行动计划请参见本报告"六、修复建议与行动计划"章节，其中包含：
• 立即行动项（0-24小时内处理{critical_failures}个严重风险和{high_failures}个高风险失败项）
• 短期改进措施（1-2周内优化低通过率战术类别）
• 中期建设规划（1-3个月建立持续验证机制）
• 长期优化方向（持续提升安全成熟度）

每项建议均包含具体实施步骤、责任团队、验证标准和预计工作量，可直接用于指导整改工作落地。
'''
        run = action_guide.add_run(guide_text)
        set_chinese_font(run)

        # 资源与支持
        doc.add_paragraph()
        resource_heading = doc.add_heading('9.1 所需资源与支持', 2)
        for run in resource_heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        resource_intro = doc.add_paragraph()
        resource_intro_text = '为确保整改工作顺利推进，需要协调以下资源和支持：'
        run = resource_intro.add_run(resource_intro_text)
        set_chinese_font(run)

        # 资源表格
        resource_table = doc.add_table(rows=6, cols=3)
        resource_table.style = 'Light Grid Accent 1'

        resource_headers = ['资源类型', '具体内容', '获取方式/来源']
        for i, header in enumerate(resource_headers):
            cell = resource_table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                set_chinese_font(run, bold=True)

        resources = [
            ('人员', '安全工程师2人、系统管理员1人、安全架构师1人', '内部调配或外部招聘'),
            ('工具', 'EDR/EPP平台、SIEM系统、漏洞扫描器、自动化测试平台', '采购或开源部署'),
            ('预算', '工具采购费、安全服务费、培训费用', '年度安全预算申请'),
            ('时间', '每月8小时维护窗口、每季度16小时测试窗口', '与运维团队协调'),
            ('授权', '系统管理员权限、安全策略变更审批、日志访问权限', '向上级申请授权')
        ]

        for idx, (res_type, content, source) in enumerate(resources, 1):
            row = resource_table.rows[idx]
            set_cell_text(row.cells[0], res_type)
            set_cell_text(row.cells[1], content)
            set_cell_text(row.cells[2], source)

        # 成功指标
        doc.add_paragraph()
        metric_heading = doc.add_heading('9.2 成功指标与验收标准', 2)
        for run in metric_heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        metrics_intro = doc.add_paragraph()
        metrics_text = '为量化整改效果，确保安全能力实质性提升，建议设定以下验收指标：'
        run = metrics_intro.add_run(metrics_text)
        set_chinese_font(run)

        # 指标表格
        metric_table = doc.add_table(rows=7, cols=4)
        metric_table.style = 'Light Grid Accent 1'

        metric_headers = ['指标名称', '当前值', '目标值', '验收标准']
        for i, header in enumerate(metric_headers):
            cell = metric_table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                set_chinese_font(run, bold=True)

        current_pass_rate = stats['pass_rate']
        current_critical_rate = (stats['risk_status']['critical']['passed'] / stats['risk_distribution']['critical'] * 100) if stats['risk_distribution']['critical'] > 0 else 100
        current_high_rate = (stats['risk_status']['high']['passed'] / stats['risk_distribution']['high'] * 100) if stats['risk_distribution']['high'] > 0 else 100

        metrics = [
            ('整体通过率', f'{current_pass_rate:.1f}%', '≥80%', '复测时达到目标值即验收通过'),
            ('严重风险通过率', f'{current_critical_rate:.1f}%', '100%', '所有严重风险测试必须全部通过'),
            ('高风险通过率', f'{current_high_rate:.1f}%', '≥95%', '高风险测试通过率不低于95%'),
            ('检测覆盖率', '-', '≥90%', '关键攻击技术检测规则覆盖率'),
            ('平均响应时间', '-', '<1小时', '高危告警从发现到响应的时间'),
            ('复测周期', '-', '每季度1次', '定期执行全面安全验证测试')
        ]

        for idx, (name, current, target, standard) in enumerate(metrics, 1):
            row = metric_table.rows[idx]
            set_cell_text(row.cells[0], name)
            set_cell_text(row.cells[1], current)
            set_cell_text(row.cells[2], target)
            set_cell_text(row.cells[3], standard)

        # 后续支持与参考资料
        doc.add_paragraph()
        contact_heading = doc.add_heading('9.3 后续支持与参考资料', 2)
        for run in contact_heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        contact_intro = doc.add_paragraph()
        contact_intro_text = '在报告解读或整改实施过程中，可参考以下官方资源获取技术支持：'
        run = contact_intro.add_run(contact_intro_text)
        set_chinese_font(run)

        # 参考资源表格
        ref_table = doc.add_table(rows=6, cols=3)
        ref_table.style = 'Light Grid Accent 1'

        ref_headers = ['资源名称', '用途说明', '访问地址']
        for i, header in enumerate(ref_headers):
            cell = ref_table.rows[0].cells[i]
            cell.text = header
            for run in cell.paragraphs[0].runs:
                set_chinese_font(run, bold=True)

        references = [
            ('MITRE ATT&CK', '完整的攻击技术定义、检测方法和缓解措施', 'https://attack.mitre.org/'),
            ('Atomic Red Team', '可执行的测试用例库和实施指南', 'https://github.com/redcanaryco/atomic-red-team'),
            ('Sigma规则仓库', '跨平台检测规则模板和最佳实践', 'https://github.com/SigmaHQ/sigma'),
            ('MITRE D3FEND', '防御技术知识库，指导安全控制部署', 'https://d3fend.mitre.org/'),
            ('MITRE ENGAGE', '主动防御策略和 adversary engagement 技术', 'https://engage.mitre.org/')
        ]

        for idx, (name, desc, url) in enumerate(references, 1):
            row = ref_table.rows[idx]
            set_cell_text(row.cells[0], name)
            set_cell_text(row.cells[1], desc)
            set_cell_text(row.cells[2], url)

        # 持续改进建议
        doc.add_paragraph()
        improve_heading = doc.add_heading('9.4 持续改进建议', 2)
        for run in improve_heading.runs:
            set_chinese_font(run, font_name='SimHei', font_size=14, bold=True)

        improve_para = doc.add_paragraph()
        improve_text = '''安全是一个持续的过程，而非一次性的任务。为确保安全能力的持续提升，建议采取以下措施：

'''
        run = improve_para.add_run(improve_text)
        set_chinese_font(run)

        improve_items = [
            ('建立常态化测试机制', '每季度执行一次完整的MITRE ATT&CK测试，持续验证安全控制的有效性'),
            ('构建安全度量体系', '建立量化指标，定期评估安全成熟度，跟踪改进趋势'),
            ('强化威胁情报能力', '订阅威胁情报源，及时更新检测规则，应对新出现的威胁技术'),
            ('开展红蓝对抗演练', '定期组织攻防演练，检验检测和响应能力，发现防御盲区'),
            ('建设安全知识库', '记录本次测试的经验教训、最佳实践，形成组织的安全知识资产'),
            ('培养安全人才队伍', '定期培训安全团队，提升对MITRE ATT&CK框架的理解和应用能力')
        ]

        for title, desc in improve_items:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(f'{title}：').bold = True
            para.add_run(desc)
            for run in para.runs:
                set_chinese_font(run)

        # 保存文档
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        doc.save(output_file)

        print(f"Word报告已生成: {output_file}")
        return output_file

def main():
    parser = argparse.ArgumentParser(description="MITRE ATT&CK 测试日志分析工具 - 针对.log文件")
    parser.add_argument("--log-dir", default="./logs", help="日志目录路径")
    parser.add_argument("--report-dir", default="./reports", help="报告输出目录")
    parser.add_argument("--format", choices=["html", "excel", "text", "docx", "all"], default="all",
                       help="报告格式: html, excel, text, docx, 或 all")
    parser.add_argument("--output", help="指定输出文件名（不含扩展名）")

    args = parser.parse_args()

    # 创建分析器
    analyzer = AtomicLogAnalyzer(args.log_dir, args.report_dir)

    print("开始分析测试日志...")
    df = analyzer.analyze_all_logs()

    if df.empty:
        print("未找到可分析的日志文件")
        sys.exit(1)

    print(f"分析完成，共处理 {len(df)} 个测试结果")

    # 生成报告
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if args.output:
        base_name = args.output
    else:
        base_name = f"security_validation_report_windows_{timestamp}"

    reports = []

    if args.format in ["html", "all"]:
        html_file = os.path.join(args.report_dir, f"{base_name}.html")
        try:
            analyzer.generate_html_report(html_file)
            reports.append(html_file)
        except Exception as e:
            print(f"生成HTML报告时出错: {e}")
            import traceback
            traceback.print_exc()

    if args.format in ["excel", "all"]:
        excel_file = os.path.join(args.report_dir, f"{base_name}.xlsx")
        try:
            analyzer.generate_excel_report(excel_file)
            reports.append(excel_file)
        except Exception as e:
            print(f"生成Excel报告时出错: {e}")
            import traceback
            traceback.print_exc()

    if args.format in ["text", "all"]:
        text_file = os.path.join(args.report_dir, f"{base_name}.txt")
        try:
            analyzer.generate_text_report(text_file)
            reports.append(text_file)
        except Exception as e:
            print(f"生成文本报告时出错: {e}")
            import traceback
            traceback.print_exc()

    if args.format in ["docx", "all"]:
        docx_file = os.path.join(args.report_dir, f"{base_name}.docx")
        try:
            result = analyzer.generate_docx_report(docx_file)
            if result:
                reports.append(docx_file)
        except Exception as e:
            print(f"生成Word报告时出错: {e}")
            import traceback
            traceback.print_exc()

    # 输出总结
    print("\n" + "=" * 60)
    print("报告生成完成!")
    for report in reports:
        print(f"- {report}")

    # 显示详细统计
    stats = analyzer.calculate_statistics()
    print("\n详细统计:")
    print(f"  总测试数: {stats['total_tests']}")
    print(f"  通过: {stats['passed']} ({stats['pass_rate']:.1f}%)")
    print(f"  失败: {stats['failed']}")
    print(f"  错误: {stats['errors']}")
    print(f"  平均执行时间: {stats['avg_execution_time']:.2f}秒")
    print(f"  测试技术: {stats['techniques_tested']}种")
    print(f"  严重风险失败: {stats['risk_status']['critical']['failed']}")
    print(f"  高风险失败: {stats['risk_status']['high']['failed']}")
    print("=" * 60)

if __name__ == "__main__":
    main()
