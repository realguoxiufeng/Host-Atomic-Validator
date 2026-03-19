#!/usr/bin/env python3
"""
主机安全验证服务方案 - Word文档生成器
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_shading(cell, fill_color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_style(doc, level, font_name='微软雅黑', font_size=16, bold=True, color=None):
    """自定义标题样式"""
    style = doc.styles[f'Heading {level}']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_document():
    doc = Document()

    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)

    # ==================== 封面 ====================
    # 添加空行
    for _ in range(6):
        doc.add_paragraph()

    # 主标题
    title = doc.add_paragraph()
    title_run = title.add_run('主机安全验证服务方案')
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.name = '微软雅黑'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('MITRE ATT&CK Based Security Validation Service')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle_run.font.name = 'Arial'
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加空行
    for _ in range(12):
        doc.add_paragraph()

    # 编制信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run('编制单位：漏洞盒子安全服务团队')
    info_run.font.size = Pt(14)
    info_run.font.name = '微软雅黑'
    info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f'编制日期：{datetime.datetime.now().strftime("%Y年%m月")}')
    date_run.font.size = Pt(14)
    date_run.font.name = '微软雅黑'
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 分页
    doc.add_page_break()

    # ==================== 目录 ====================
    doc.add_heading('目录', level=1)

    toc_items = [
        ('一、服务概述', 1),
        ('    1.1 服务背景', 2),
        ('    1.2 服务目标', 2),
        ('    1.3 服务范围', 2),
        ('二、技术框架', 1),
        ('    2.1 MITRE ATT&CK 框架', 2),
        ('    2.2 测试工具组件', 2),
        ('三、服务内容', 1),
        ('    3.1 Linux 主机安全验证', 2),
        ('    3.2 Windows 主机安全验证', 2),
        ('    3.3 容器安全验证', 2),
        ('四、服务实施流程', 1),
        ('    4.1 实施阶段划分', 2),
        ('    4.2 详细实施步骤', 2),
        ('    4.3 执行检查清单', 2),
        ('五、交付成果', 1),
        ('六、风险控制', 1),
        ('七、服务报价', 1),
        ('八、成功案例', 1),
        ('九、附录', 1),
    ]

    for item, level in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.paragraph_format.line_spacing = 1.5

    doc.add_page_break()

    # ==================== 正文 ====================

    # 一、服务概述
    doc.add_heading('一、服务概述', level=1)

    doc.add_heading('1.1 服务背景', level=2)
    doc.add_paragraph(
        '随着网络安全威胁日益复杂，传统的安全防护措施面临巨大挑战。主机作为企业核心业务承载平台，'
        '其安全性直接关系到业务连续性和数据安全。本服务基于 MITRE ATT&CK 框架，采用 Atomic Red Team '
        '方法论，通过模拟真实攻击行为，验证主机安全防护体系的有效性。'
    )

    doc.add_heading('1.2 服务目标', level=2)
    goals = [
        '验证现有安全防护措施的有效性',
        '发现主机安全盲点和防护薄弱环节',
        '评估安全监控和告警机制',
        '提供可落地的安全加固建议',
        '输出专业的安全验证报告'
    ]
    for goal in goals:
        p = doc.add_paragraph(goal, style='List Bullet')

    doc.add_heading('1.3 服务范围', level=2)

    # 服务范围表格
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['平台', '测试技术数量', '覆盖战术']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].font.bold = True

    data = [
        ('Linux', '42项', '8大战术类别'),
        ('Windows', '114项', '12大战术类别'),
        ('容器环境', '6项', 'Kubernetes安全')
    ]
    for i, row_data in enumerate(data, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    # 二、技术框架
    doc.add_heading('二、技术框架', level=1)

    doc.add_heading('2.1 MITRE ATT&CK 框架', level=2)
    doc.add_paragraph('本服务基于 MITRE ATT&CK 框架进行安全验证，覆盖以下战术类别：')

    # 战术表格
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'

    headers = ['战术类别', '英文名称', '验证重点']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].font.bold = True

    tactics = [
        ('初始访问', 'Initial Access', '外部入侵途径验证'),
        ('执行', 'Execution', '恶意代码执行检测'),
        ('持久化', 'Persistence', '长期驻留机制检测'),
        ('权限提升', 'Privilege Escalation', '提权漏洞验证'),
        ('防御规避', 'Defense Evasion', '安全规避能力测试'),
        ('凭据访问', 'Credential Access', '凭据窃取防护验证'),
        ('发现', 'Discovery', '信息收集检测能力'),
        ('横向移动', 'Lateral Movement', '内网穿透防护')
    ]
    for i, row_data in enumerate(tactics, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    doc.add_heading('2.2 测试工具组件', level=2)

    # 目录结构
    dir_structure = '''/opt/bas/
├── goart-linux              # 测试执行引擎
├── run_all_atomics.sh       # 批量测试脚本
├── analyze_atomic_logs2.py  # 日志分析工具
├── configs/                 # 配置文件目录
│   ├── execution_plan.json  # 执行计划配置
│   ├── test_linux.txt       # Linux测试列表
│   └── test_windos.txt      # Windows测试列表
├── logs/                    # 测试日志目录
└── reports/                 # 报告输出目录'''

    p = doc.add_paragraph()
    run = p.add_run(dir_structure)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # 三、服务内容
    doc.add_heading('三、服务内容', level=1)

    doc.add_heading('3.1 Linux 主机安全验证', level=2)

    doc.add_heading('3.1.1 核心测试项目', level=3)

    # Linux测试表格
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Table Grid'

    headers = ['技术ID', '技术名称', '风险等级', '验证目标']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].font.bold = True

    linux_tests = [
        ('T1003.007', 'Proc文件系统凭据转储', 'Critical', '检测凭据窃取行为'),
        ('T1003.008', '/etc/passwd/shadow文件读取', 'Critical', '敏感文件访问监控'),
        ('T1053.003', 'Cron计划任务', 'High', '持久化检测能力'),
        ('T1543.002', 'Systemd服务创建', 'High', '服务创建监控'),
        ('T1547.006', '内核模块加载', 'Critical', 'Rootkit检测'),
        ('T1548.001', 'Setuid/Setgid权限滥用', 'High', '提权行为检测'),
        ('T1548.003', 'Sudo权限滥用', 'High', 'Sudo滥用监控'),
        ('T1070.002', '系统日志清除', 'Medium', '日志篡改检测'),
        ('T1556.003', 'PAM模块篡改', 'Critical', '认证机制保护'),
        ('T1574.006', 'LD_PRELOAD劫持', 'High', '动态库劫持检测')
    ]
    for i, row_data in enumerate(linux_tests, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    doc.add_heading('3.1.2 重点验证场景', level=3)

    # 场景一
    p = doc.add_paragraph()
    run = p.add_run('场景一：凭据窃取防护验证')
    run.font.bold = True

    doc.add_paragraph('测试技术：T1003.007, T1003.008, T1552.003')
    doc.add_paragraph('验证目标：')
    for item in ['是否能检测到/etc/shadow文件读取', '是否能检测到/proc/*/mem敏感访问', '是否能发现bash历史中的敏感信息']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph('预期结果：安全设备应产生告警')

    # 场景二
    p = doc.add_paragraph()
    run = p.add_run('场景二：持久化机制检测')
    run.font.bold = True

    doc.add_paragraph('测试技术：T1053.003, T1543.002, T1546.004, T1546.005')
    doc.add_paragraph('验证目标：')
    for item in ['Cron任务创建是否被监控', 'Systemd服务创建是否告警', 'Shell配置文件修改是否检测']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph('预期结果：持久化行为应被记录和告警')

    doc.add_heading('3.2 Windows 主机安全验证', level=2)

    doc.add_heading('3.2.1 核心测试项目', level=3)

    # Windows测试表格
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].font.bold = True

    windows_tests = [
        ('T1003.001', 'LSASS内存凭据转储', 'Critical', '凭据保护能力'),
        ('T1003.002', 'SAM数据库转储', 'Critical', '本地凭据保护'),
        ('T1055', '进程注入', 'Critical', '进程保护机制'),
        ('T1059.001', 'PowerShell执行', 'High', '脚本执行监控'),
        ('T1053.005', '计划任务创建', 'High', '持久化检测'),
        ('T1547.001', '注册表启动项', 'High', '启动项监控'),
        ('T1070.001', 'Windows事件日志清除', 'High', '日志保护验证'),
        ('T1558.003', 'Kerberoasting攻击', 'Critical', '域安全验证'),
        ('T1550.002', 'Pass-the-Hash攻击', 'Critical', '认证安全验证'),
        ('T1218', '签名二进制代理执行', 'High', 'LOLBins检测')
    ]
    for i, row_data in enumerate(windows_tests, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    doc.add_heading('3.3 容器安全验证', level=2)

    # 容器测试表格
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].font.bold = True

    container_tests = [
        ('T1609', 'Kubernetes Exec容器执行', 'Critical', '容器执行监控'),
        ('T1610', '容器部署', 'High', '容器创建监控'),
        ('T1611', '容器逃逸到主机', 'Critical', '容器隔离验证'),
        ('T1612', '主机上构建镜像', 'Medium', '镜像构建监控'),
        ('T1613', '容器和资源发现', 'Low', '信息收集检测'),
        ('T1552.007', 'Kubernetes Secrets获取', 'Critical', '密钥保护验证')
    ]
    for i, row_data in enumerate(container_tests, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    # 四、服务实施流程
    doc.add_heading('四、服务实施流程', level=1)

    doc.add_heading('4.1 实施阶段划分', level=2)

    # 阶段表格
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    phase_headers = ['阶段', '工作内容', '预计时间']
    for i, header in enumerate(phase_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].font.bold = True

    phases = [
        ('准备阶段', '环境准备、工具部署', '1-2天'),
        ('执行阶段', '安全测试执行', '2-3天'),
        ('分析阶段', '日志分析、结果评估', '1-2天'),
        ('报告阶段', '报告编写、成果汇报', '1天')
    ]
    for i, row_data in enumerate(phases, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    doc.add_heading('4.2 详细实施步骤', level=2)

    # 步骤一
    p = doc.add_paragraph()
    run = p.add_run('第一步：环境准备')
    run.font.bold = True

    cmd1 = '''# 部署测试工具
mkdir -p /opt/bas/{logs,reports,configs}
cd /opt/bas

# 安装Python依赖
pip install -r requirements.txt

# 克隆Atomic Red Team项目
git clone https://github.com/redcanaryco/atomic-red-team.git /opt/atomic-red-team

# 设置执行权限
chmod +x /opt/bas/goart-linux
chmod +x /opt/bas/run_all_atomics.sh'''

    p = doc.add_paragraph()
    run = p.add_run(cmd1)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # 步骤二
    p = doc.add_paragraph()
    run = p.add_run('第二步：执行测试')
    run.font.bold = True

    cmd2 = '''# Linux主机测试
./run_all_atomics.sh -l configs/test_linux.txt -s -v

# Windows主机测试
./run_all_atomics.sh -l configs/test_windos.txt -s -v

# 容器安全测试
./run_all_atomics.sh -t T1609 -v
./run_all_atomics.sh -t T1611 -v'''

    p = doc.add_paragraph()
    run = p.add_run(cmd2)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # 步骤三
    p = doc.add_paragraph()
    run = p.add_run('第三步：日志分析与报告生成')
    run.font.bold = True

    cmd3 = '''# 生成全部格式报告
python3 analyze_atomic_logs2.py

# 仅生成HTML报告
python3 analyze_atomic_logs2.py --format html'''

    p = doc.add_paragraph()
    run = p.add_run(cmd3)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading('4.3 执行检查清单', level=2)

    p = doc.add_paragraph()
    run = p.add_run('测试前检查：')
    run.font.bold = True

    checklist_before = [
        '确认测试环境已与生产环境隔离',
        '获得测试授权书和测试时间窗口',
        '确认安全设备已开启日志记录',
        '备份关键系统配置',
        '确认回滚方案已准备就绪'
    ]
    for item in checklist_before:
        doc.add_paragraph('□ ' + item)

    p = doc.add_paragraph()
    run = p.add_run('测试后检查：')
    run.font.bold = True

    checklist_after = [
        '确认所有测试已完成',
        '检查系统状态是否正常',
        '收集全部日志文件',
        '执行系统恢复操作（如需要）'
    ]
    for item in checklist_after:
        doc.add_paragraph('□ ' + item)

    # 五、交付成果
    doc.add_heading('五、交付成果', level=1)

    doc.add_heading('5.1 报告内容', level=2)

    report_sections = [
        ('执行摘要', '测试概况和关键发现、高风险问题汇总、整体安全状况评估'),
        ('详细测试结果', '测试统计、战术覆盖、风险分析、技术详情'),
        ('告警评估', '安全设备告警效果分析、防护盲点识别'),
        ('加固建议', '按优先级分类的安全加固建议')
    ]

    for title, content in report_sections:
        p = doc.add_paragraph()
        run = p.add_run(title + '：')
        run.font.bold = True
        p.add_run(content)

    doc.add_heading('5.2 交付物清单', level=2)

    # 交付物表格
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'

    headers = ['序号', '交付物', '格式', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].font.bold = True

    deliverables = [
        ('1', '安全验证报告', 'HTML/DOCX/PDF', '完整测试报告'),
        ('2', '测试数据汇总', 'XLSX', 'Excel格式详细数据'),
        ('3', '原始日志文件', 'LOG', '测试执行原始日志'),
        ('4', '测试汇总文件', 'TXT', '文本格式汇总'),
        ('5', '加固建议清单', 'DOCX', '可执行的安全建议'),
        ('6', '测试配置文件', 'JSON/TXT', '可复用的测试配置')
    ]
    for i, row_data in enumerate(deliverables, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    # 六、风险控制
    doc.add_heading('六、风险控制', level=1)

    doc.add_heading('6.1 风险识别', level=2)

    # 风险表格
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['风险类型', '风险描述', '影响程度']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].font.bold = True

    risks = [
        ('业务影响', '测试可能导致服务中断', '高'),
        ('数据安全', '测试可能涉及敏感数据', '高'),
        ('误报风险', '安全设备可能产生大量告警', '中'),
        ('残留痕迹', '测试可能留下持久化痕迹', '中')
    ]
    for i, row_data in enumerate(risks, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    doc.add_heading('6.2 风险控制措施', level=2)

    measures = [
        ('环境隔离', '在独立的测试环境中执行测试，使用快照功能便于快速恢复'),
        ('时间窗口', '选择业务低峰期执行测试，预留足够的测试和恢复时间'),
        ('数据保护', '测试前备份关键数据，避免在生产数据上执行测试'),
        ('回滚机制', '记录所有配置变更，准备系统快照，制定详细的回滚步骤')
    ]

    for title, content in measures:
        p = doc.add_paragraph()
        run = p.add_run('• ' + title + '：')
        run.font.bold = True
        p.add_run(content)

    # 七、服务报价
    doc.add_heading('七、服务报价', level=1)

    doc.add_heading('7.1 服务周期', level=2)

    doc.add_paragraph('预计服务周期：5-8个工作日')

    doc.add_heading('7.2 服务报价说明', level=2)

    # 报价表格
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['项目', '内容', '备注']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].font.bold = True

    pricing = [
        ('基础服务费', 'Linux/Windows主机安全验证', '含报告输出'),
        ('加急服务', '3天内完成全部工作', '加收30%'),
        ('复测服务', '加固后复测验证', '优惠50%'),
        ('咨询服务', '安全加固技术支持', '按人天计费')
    ]
    for i, row_data in enumerate(pricing, 1):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text

    doc.add_paragraph()

    doc.add_heading('7.3 服务前提', level=2)

    prerequisites = [
        '客户提供独立的测试环境',
        '客户提供测试授权书',
        '客户配合提供必要的环境信息',
        '客户安排技术人员配合测试'
    ]
    for item in prerequisites:
        doc.add_paragraph(item, style='List Bullet')

    # 八、成功案例
    doc.add_heading('八、成功案例', level=1)

    doc.add_heading('8.1 某金融机构主机安全验证', level=2)

    p = doc.add_paragraph()
    run = p.add_run('项目背景：')
    run.font.bold = True
    doc.add_paragraph('客户：某省级银行')
    doc.add_paragraph('范围：核心业务系统服务器（Linux 50台，Windows 30台）')

    p = doc.add_paragraph()
    run = p.add_run('测试成果：')
    run.font.bold = True
    doc.add_paragraph('执行测试技术：156项')
    doc.add_paragraph('发现防护盲点：23个')
    doc.add_paragraph('Critical级别问题：5个')
    doc.add_paragraph('High级别问题：12个')

    p = doc.add_paragraph()
    run = p.add_run('改进效果：')
    run.font.bold = True
    doc.add_paragraph('部署新检测规则：47条')
    doc.add_paragraph('防护覆盖率提升：65% → 92%')
    doc.add_paragraph('告警准确率提升：78% → 95%')

    doc.add_heading('8.2 某互联网企业容器安全验证', level=2)

    p = doc.add_paragraph()
    run = p.add_run('项目背景：')
    run.font.bold = True
    doc.add_paragraph('客户：某电商公司')
    doc.add_paragraph('范围：Kubernetes集群（3个集群，200+节点）')

    p = doc.add_paragraph()
    run = p.add_run('关键发现：')
    run.font.bold = True
    findings = ['容器逃逸检测能力缺失', 'Kubernetes Secrets保护不足', '特权容器管控不严格']
    for item in findings:
        doc.add_paragraph(item, style='List Bullet')

    # 九、附录
    doc.add_heading('九、附录', level=1)

    doc.add_heading('附录A：测试技术完整清单', level=2)
    doc.add_paragraph('详见配置文件：')
    doc.add_paragraph('• /opt/bas/configs/test_linux.txt - Linux测试清单（42项）')
    doc.add_paragraph('• /opt/bas/configs/test_windos.txt - Windows测试清单（114项）')

    doc.add_heading('附录B：常用命令速查', level=2)

    cmd_ref = '''# 查看帮助
./run_all_atomics.sh -h

# 干运行（预览）
./run_all_atomics.sh -l configs/test_linux.txt -d

# 执行测试并跳过失败
./run_all_atomics.sh -l configs/test_linux.txt -s

# 分析日志生成报告
python3 analyze_atomic_logs2.py --format all

# 仅生成HTML报告
python3 analyze_atomic_logs2.py --format html'''

    p = doc.add_paragraph()
    run = p.add_run(cmd_ref)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # 分页后添加末尾信息
    doc.add_page_break()

    # 编制信息
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—— 文档结束 ——')
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('编制单位：漏洞盒子安全服务团队')
    run.font.size = Pt(12)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'编制日期：{datetime.datetime.now().strftime("%Y年%m月")}')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('文档版本：v1.0')
    run.font.size = Pt(12)

    return doc

def main():
    print("正在生成主机安全验证服务方案文档...")
    doc = create_document()
    output_file = "主机安全验证服务方案.docx"
    doc.save(output_file)
    print(f"文档已生成: {output_file}")

if __name__ == "__main__":
    main()